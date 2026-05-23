#!/usr/bin/env python3
"""Draft ALTA Title Commitment (2021) for 502 Ivy Green Lane, Nexton (Lot 1635)."""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Draft Commitment - 502 Ivy Green Lane Summerville.docx")


def add_page_break(doc): doc.add_page_break()


def set_narrow_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.75); s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)


def add_header_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "000000")
    pBdr.append(b); pPr.append(pBdr)


def add_right_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.font.size = Pt(9)


def add_page_header(doc):
    add_right_header(doc, "Commitment for Title Insurance")
    add_right_header(doc, "South Carolina - 2021 v. 01.00 (07-01-2021)")
    add_header_line(doc)


def add_footer_text(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    t = OxmlElement("w:top")
    t.set(qn("w:val"), "single"); t.set(qn("w:sz"), "4")
    t.set(qn("w:space"), "1"); t.set(qn("w:color"), "999999")
    pBdr.append(t); pPr.append(pBdr)
    r = p.add_run(
        "This page is only a part of a 2021 ALTA Commitment for Title Insurance. "
        "This Commitment is not valid without the Notice; the Commitment to Issue Policy; "
        "the Commitment Conditions; Schedule A; Schedule B, Part I\u2014Requirements; "
        "Schedule B, Part II\u2014Exceptions; and a counter-signature by the Company or "
        "its issuing agent that may be in electronic form."
    )
    r.font.size = Pt(7); r.font.italic = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(0)
    pPr2 = p2._p.get_or_add_pPr(); pBdr2 = OxmlElement("w:pBdr")
    t2 = OxmlElement("w:top")
    t2.set(qn("w:val"), "single"); t2.set(qn("w:sz"), "4")
    t2.set(qn("w:space"), "1"); t2.set(qn("w:color"), "000000")
    pBdr2.append(t2); pPr2.append(pBdr2)
    r2 = p2.add_run(
        "Copyright 2021 American Land Title Association. All rights reserved.\n"
        "The use of this Form (or any derivative thereof) is restricted to ALTA licensees and\n"
        "ALTA members in good standing as of the date of use. All other uses are prohibited.\n"
        "Reprinted under license from the American Land Title Association."
    )
    r2.font.size = Pt(7); r2.font.bold = True


def p_normal(doc, text, bold=False, size=10, space_after=6, space_before=0, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment: p.alignment = alignment
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    return p


def p_center_bold(doc, text, size=12, space_before=12, space_after=8):
    return p_normal(doc, text, bold=True, size=size, space_after=space_after,
                    space_before=space_before, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_numbered_item(doc, number, text, bold_number=True, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    rn = p.add_run(f"{number}.\t"); rn.font.size = Pt(10); rn.font.bold = bold_number
    rt = p.add_run(text); rt.font.size = Pt(10)
    return p


def add_sub_item(doc, letter, text, indent=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    rl = p.add_run(f"{letter}.\t"); rl.font.size = Pt(10)
    rt = p.add_run(text); rt.font.size = Pt(10)
    return p


def add_group_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text); r.font.size = Pt(10); r.font.bold = True; r.font.italic = True


# ============================================================
#  BUILD
# ============================================================
doc = Document()
set_narrow_margins(doc)
style = doc.styles["Normal"]; style.font.name = "Times New Roman"; style.font.size = Pt(10)

# ---------- PAGE 1 — NOTICE / COMMITMENT TO ISSUE POLICY ----------
add_page_header(doc)

p_center_bold(doc, "ALTA COMMITMENT FOR TITLE INSURANCE\nissued by\n[INSURANCE CARRIER]", size=11, space_before=24)
p_center_bold(doc, "NOTICE", size=11, space_before=18)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
r = p.add_run("IMPORTANT\u2014READ CAREFULLY"); r.font.size = Pt(10); r.font.bold = True
r2 = p.add_run(
    ": THIS COMMITMENT IS AN OFFER TO ISSUE ONE OR MORE TITLE INSURANCE POLICIES. "
    "ALL CLAIMS OR REMEDIES SOUGHT AGAINST THE COMPANY INVOLVING THE CONTENT OF THIS "
    "COMMITMENT OR THE POLICY MUST BE BASED SOLELY IN CONTRACT."
); r2.font.size = Pt(10)

p_normal(doc,
    "THIS COMMITMENT IS NOT AN ABSTRACT OF TITLE, REPORT OF THE CONDITION OF TITLE, LEGAL "
    "OPINION, OPINION OF TITLE, OR OTHER REPRESENTATION OF THE STATUS OF TITLE. THE "
    "PROCEDURES USED BY THE COMPANY TO DETERMINE INSURABILITY OF THE TITLE, INCLUDING ANY "
    "SEARCH AND EXAMINATION, ARE PROPRIETARY TO THE COMPANY, WERE PERFORMED SOLELY FOR "
    "THE BENEFIT OF THE COMPANY, AND CREATE NO EXTRACONTRACTUAL LIABILITY TO ANY PERSON, "
    "INCLUDING A PROPOSED INSURED.", size=10, space_after=8)

p_normal(doc,
    "THE COMPANY\u2019S OBLIGATION UNDER THIS COMMITMENT IS TO ISSUE A POLICY TO A PROPOSED "
    "INSURED IDENTIFIED IN SCHEDULE A IN ACCORDANCE WITH THE TERMS AND CONDITIONS OF THIS "
    "COMMITMENT. THE COMPANY HAS NO LIABILITY OR OBLIGATION INVOLVING THE CONTENT OF THIS "
    "COMMITMENT TO ANY OTHER PERSON.", size=10, space_after=12)

p_center_bold(doc, "COMMITMENT TO ISSUE POLICY", size=11)

p_normal(doc,
    'Subject to the Notice; Schedule B, Part I\u2014Requirements; Schedule B, Part II\u2014Exceptions; '
    'and the Commitment Conditions, [Insurance Carrier], a [State] Corporation (the \u201cCompany\u201d), '
    'commits to issue the Policy according to the terms and provisions of this Commitment. This '
    'Commitment is effective as of the Commitment Date shown in Schedule A for each Policy described '
    'in Schedule A, only when the Company has entered in Schedule A both the specified dollar amount '
    'as the Proposed Amount of Insurance and the name of the Proposed Insured.', space_after=8)

p_normal(doc,
    "If all of the Schedule B, Part I\u2014Requirements have not been met within 180 days after the "
    "Commitment Date, this Commitment terminates and the Company\u2019s liability and obligation end.",
    space_after=12)

p_normal(doc, "[INSURANCE CARRIER]", bold=True, space_after=18)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
r = p.add_run("By: ____________________________\t\t\tBy: ____________________________"); r.font.size = Pt(10)

p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("      [Name], President\t\t\t\t\t      [Name], Secretary")
r2.font.size = Pt(9); r2.font.bold = True

add_footer_text(doc)

# ---------- COMMITMENT CONDITIONS ----------
add_page_break(doc); add_page_header(doc)
p_center_bold(doc, "COMMITMENT CONDITIONS", size=11, space_before=6)

add_numbered_item(doc, 1, "DEFINITIONS")
for letter, text in [
    ("a", '"Discriminatory Covenant": Any covenant, condition, restriction, or limitation that is unenforceable under applicable law because it illegally discriminates against a class of individuals based on personal characteristics such as race, color, religion, sex, sexual orientation, gender identity, familial status, disability, national origin, or other legally protected class.'),
    ("b", '"Knowledge" or "Known": Actual knowledge or actual notice, but not constructive notice imparted by the Public Records.'),
    ("c", '"Land": The land described in Item 5 of Schedule A and improvements located on that land that by State law constitute real property.'),
    ("d", '"Mortgage": A mortgage, deed of trust, trust deed, security deed, or other real property security instrument, including one evidenced by electronic means authorized by law.'),
    ("e", '"Policy": Each contract of title insurance, in a form adopted by the American Land Title Association, issued or to be issued by the Company pursuant to this Commitment.'),
    ("f", '"Proposed Amount of Insurance": Each dollar amount specified in Schedule A as the Proposed Amount of Insurance of each Policy to be issued pursuant to this Commitment.'),
    ("g", '"Proposed Insured": Each person identified in Schedule A as the Proposed Insured of each Policy to be issued pursuant to this Commitment.'),
    ("h", '"Public Records": The recording or filing system established under State statutes in effect at the Commitment Date under which a document must be recorded or filed to impart constructive notice of matters relating to the Title to a purchaser for value without Knowledge.'),
    ("i", '"State": The state or commonwealth of the United States within whose exterior boundaries the Land is located.'),
    ("j", '"Title": The estate or interest in the Land identified in Item 3 of Schedule A.'),
]:
    add_sub_item(doc, letter, text)

add_numbered_item(doc, 2,
    "If all of the Schedule B, Part I\u2014Requirements have not been met within the time period specified "
    "in the Commitment to Issue Policy, this Commitment terminates and the Company\u2019s liability and obligation end.")

add_numbered_item(doc, 3, "The Company\u2019s liability and obligation is limited by and this Commitment is not valid without:")
for letter, text in [("a", "the Notice;"), ("b", "the Commitment to Issue Policy;"),
                      ("c", "the Commitment Conditions;"), ("d", "Schedule A;"),
                      ("e", "Schedule B, Part I\u2014Requirements;"),
                      ("f", "Schedule B, Part II\u2014Exceptions; and"),
                      ("g", "a counter-signature by the Company or its issuing agent that may be in electronic form.")]:
    add_sub_item(doc, letter, text)

add_page_break(doc); add_page_header(doc)

add_numbered_item(doc, 4, "COMPANY\u2019S RIGHT TO AMEND")
p_normal(doc,
    "The Company may amend this Commitment at any time. If the Company amends this Commitment to add "
    "a defect, lien, encumbrance, adverse claim, or other matter recorded in the Public Records prior to the "
    "Commitment Date, any liability of the Company is limited by Commitment Condition 5. The Company is not "
    "liable for any other amendment to this Commitment.", space_after=8)

add_numbered_item(doc, 5, "LIMITATIONS OF LIABILITY")
add_sub_item(doc, "a",
    "The Company\u2019s liability under Commitment Condition 4 is limited to the Proposed Insured\u2019s actual "
    "expense incurred in the interval between the Company\u2019s delivery to the Proposed Insured of the "
    "Commitment and the delivery of the amended Commitment, resulting from the Proposed Insured\u2019s good "
    "faith reliance to:")
for sub in ["comply with the Schedule B, Part I\u2014Requirements;",
            "eliminate, with the Company\u2019s written consent, any Schedule B, Part II\u2014Exceptions; or",
            "acquire the Title or create the Mortgage covered by this Commitment."]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(1.5)
    r = p.add_run(sub); r.font.size = Pt(10)

add_sub_item(doc, "b", "The Company is not liable under Commitment Condition 5.a. if the Proposed Insured requested the amendment or had Knowledge of the matter and did not notify the Company about it in writing.")
add_sub_item(doc, "c", "The Company is only liable under Commitment Condition 4 if the Proposed Insured would not have incurred the expense had the Commitment included the added matter when the Commitment was first delivered to the Proposed Insured.")
add_sub_item(doc, "d", "The Company\u2019s liability does not exceed the lesser of the Proposed Insured\u2019s actual expense incurred in good faith and described in Commitment Condition 5.a. or the Proposed Amount of Insurance.")
add_sub_item(doc, "e", "The Company is not liable for the content of the Transaction Identification Data, if any.")
add_sub_item(doc, "f", "The Company is not obligated to issue the Policy referred to in this Commitment unless all of the Schedule B, Part I\u2014Requirements have been met to the satisfaction of the Company.")
add_sub_item(doc, "g", "The Company\u2019s liability is further limited by the terms and provisions of the Policy to be issued to the Proposed Insured.")

add_numbered_item(doc, 6, "LIABILITY OF THE COMPANY MUST BE BASED ON THIS COMMITMENT; CHOICE OF LAW AND CHOICE OF FORUM")
add_sub_item(doc, "a", "Only a Proposed Insured identified in Schedule A, and no other person, may make a claim under this Commitment.")
add_sub_item(doc, "b", "Any claim must be based in contract under the State law of the State where the Land is located and is restricted to the terms and provisions of this Commitment.")
add_sub_item(doc, "c", "This Commitment, as last revised, is the exclusive and entire agreement between the parties with respect to the subject matter of this Commitment and supersedes all prior commitment negotiations, representations, and proposals.")
add_sub_item(doc, "d", "The deletion or modification of any Schedule B, Part II\u2014Exception does not constitute an agreement or obligation to provide coverage beyond the terms and provisions of this Commitment or the Policy.")
add_sub_item(doc, "e", "Any amendment or endorsement to this Commitment must be in writing and authenticated by a person authorized by the Company.")
add_sub_item(doc, "f", "When the Policy is issued, all liability and obligation under this Commitment will end and the Company\u2019s only liability will be under the Policy.")

add_page_break(doc); add_page_header(doc)

add_numbered_item(doc, 7, "IF THIS COMMITMENT IS ISSUED BY AN ISSUING AGENT")
p_normal(doc, "The issuing agent is the Company\u2019s agent only for the limited purpose of issuing title insurance commitments and policies. The issuing agent is not the Company\u2019s agent for closing, settlement, escrow, or any other purpose.", space_after=8)

add_numbered_item(doc, 8, "PRO-FORMA POLICY")
p_normal(doc, "The Company may provide, at the request of a Proposed Insured, a pro-forma policy illustrating the coverage that the Company may provide. A pro-forma policy neither reflects the status of Title at the time that the pro-forma policy is delivered to a Proposed Insured, nor is it a commitment to insure.", space_after=8)

add_numbered_item(doc, 9, "CLAIMS PROCEDURES")
p_normal(doc, "This Commitment incorporates by reference all Conditions for making a claim in the Policy to be issued to the Proposed Insured. Commitment Condition 9 does not modify the limitations of liability in Commitment Conditions 5 and 6.", space_after=8)

add_numbered_item(doc, 10, "CLASS ACTION")
p_normal(doc, "ALL CLAIMS AND DISPUTES ARISING OUT OF OR RELATING TO THIS COMMITMENT, INCLUDING ANY SERVICE OR OTHER MATTER IN CONNECTION WITH ISSUING THIS COMMITMENT, ANY BREACH OF A COMMITMENT PROVISION, OR ANY OTHER CLAIM OR DISPUTE ARISING OUT OF OR RELATING TO THE TRANSACTION GIVING RISE TO THIS COMMITMENT, MUST BE BROUGHT IN AN INDIVIDUAL CAPACITY. NO PARTY MAY SERVE AS PLAINTIFF, CLASS MEMBER, OR PARTICIPANT IN ANY CLASS OR REPRESENTATIVE PROCEEDING. ANY POLICY ISSUED PURSUANT TO THIS COMMITMENT WILL CONTAIN A CLASS ACTION CONDITION.", size=10, space_after=8)

add_numbered_item(doc, 11, "ARBITRATION")
p_normal(doc, "The Policy contains an arbitration clause. All arbitrable matters when the Proposed Amount of Insurance is $2,000,000 or less may be arbitrated at the election of either the Company or the Proposed Insured as the exclusive remedy of the parties. A Proposed Insured may review a copy of the arbitration rules at http://www.alta.org/arbitration.", space_after=8)

add_footer_text(doc)

# ---------- SCHEDULE A ----------
add_page_break(doc); add_page_header(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
r = p.add_run("Transaction Identification Data, for which the Company assumes no liability as set forth in Commitment Condition 5.e.:")
r.font.size = Pt(10); r.font.bold = True

for label, value in [
    ("Issuing Agent:", ""),
    ("Issuing Office:", ""),
    ("", ""),
    ("Issuing Office\u2019s ALTA\u00ae Registry ID:", ""),
    ("Loan ID Number:", ""),
    ("Commitment Number:", ""),
    ("Issuing Office File Number:", ""),
    ("Property Address:", "502 Ivy Green Lane, Summerville, SC 29486"),
    ("Revision Number:", ""),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    if label:
        rl = p.add_run(f"{label} "); rl.font.size = Pt(10)
    rv = p.add_run(value); rv.font.size = Pt(10)

p_center_bold(doc, "SCHEDULE A", size=12, space_before=18)

add_numbered_item(doc, 1, "Commitment Date:  at 8:00 AM")

add_numbered_item(doc, 2, "Policy to be issued:")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = Inches(1.0)
for left in [
    "a.\t2021 ALTA Owner\u2019s Policy",
    "\tProposed Insured:\t\t[Proposed Insured]",
    "\tProposed Amount of Insurance:\t$___________",
    "\tThe estate or interest to be insured:\tfee simple",
]:
    r = p.add_run(f"{left}\n"); r.font.size = Pt(10)

add_numbered_item(doc, 3, "The estate or interest in the Land at the Commitment Date is:\n\tfee simple")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.first_line_indent = Inches(-0.35)
rn = p.add_run("4.\t"); rn.font.size = Pt(10); rn.font.bold = True
rt = p.add_run(
    "The Title is, at the Commitment Date, vested in:\n"
    "\tAshton Charleston Residential L.L.C., a South Carolina limited liability company, "
    "by Limited Warranty Deed from NASH \u2013 Nexton Holdings, LLC, a Delaware limited liability "
    "company, dated June 11, 2025, and recorded June 13, 2025, with the Berkeley County Register "
    "of Deeds in Book 5229, Page 36 (Instrument No. 2025019317)."
); rt.font.size = Pt(10)

add_numbered_item(doc, 5, "The Land is described as follows:")
p_normal(doc, "SEE EXHIBIT A ATTACHED HERETO AND MADE A PART HEREOF", bold=True, space_before=6, space_after=18)

p_normal(doc, "[TITLE AGENCY NAME]", bold=True, space_after=2)
p_normal(doc, "[Address]\n[City, State ZIP]\nTelephone:", space_after=12)
p_normal(doc, "Countersigned:", space_after=12)

p = doc.add_paragraph(); r = p.add_run("By:____________________________"); r.font.size = Pt(10)
p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("      Authorized Signatory"); r2.font.size = Pt(10); r2.font.bold = True
p3 = doc.add_paragraph()
r3 = p3.add_run("[Name], License #[______]\n[Title Agency], License #"); r3.font.size = Pt(10)

add_footer_text(doc)

# ---------- EXHIBIT A ----------
add_page_break(doc); add_page_header(doc)
p_center_bold(doc, "EXHIBIT A", size=12, space_before=6)
p_normal(doc, "Commitment No.: [File Number]", bold=True, space_after=12)

p_normal(doc,
    "The land referred to herein below is situated in the County of Berkeley, State of South Carolina, "
    "and is described as follows:", space_after=12)

p_normal(doc,
    "ALL those certain pieces, parcels or lots of land, situate, lying and being in the County of Berkeley, "
    "State of South Carolina, shown and designated as \u201cLOT 1635\u201d on a plat entitled \u201cSUBDIVISION PLAT "
    "MIDTOWN \u2013 PHASE 3C-2, HOA AREA NOS. 81\u201383 AND RESIDUAL \u2018TRACT 2A\u2019 OF NEXTON SHOWING THE "
    "SUBDIVISION OF TRACT 2A, TMS NO. 195-00-00-124 TO CREATE MIDTOWN \u2013 PHASE 3C-2 (85 LOTS), "
    "IVY GREEN LANE, JUNE BERRY DRIVE, TWINFLOWER LANE, WOOD LILLY LOOP, HOA AREA NOS. 81\u201383, "
    "PRIVATE R/W NO. 74, NEW BCWS UTILITY EASEMENTS, NEW DRAINAGE EASEMENTS, NEW BEC EASEMENTS, "
    "AND RESIDUAL \u2018TRACT 2A\u2019 OWNED BY NASH \u2013 NEXTON HOLDINGS, LLC LOCATED IN BERKELEY COUNTY, "
    "SOUTH CAROLINA,\u201d prepared by Johnathan F. Burns, PLS #22742, of GPA, Inc., dated November 14, "
    "2024, and recorded April 3, 2025, at Instrument Nos. 2025010343\u20132025010346, in the Office of the "
    "Register of Deeds for Berkeley County, South Carolina, reference to which is hereby craved for a "
    "more complete description.", space_after=8)

p_normal(doc, "TMS No.: 209-14-02-015 (Lot 1635)", bold=True, space_after=8)

p_normal(doc,
    "Street Address: 502 Ivy Green Lane, Summerville, SC 29486.", space_after=8)

p_normal(doc,
    "This being a portion of the property conveyed to NASH \u2013 Nexton Holdings, LLC by deed of "
    "MWV-Parks of Berkeley, LLC dated March 3, 2017, and recorded March 3, 2017, in Book 2406, "
    "Page 303, in the Office of the Register of Deeds for Berkeley County, South Carolina; and "
    "being the same property conveyed to Ashton Charleston Residential L.L.C. by Limited Warranty "
    "Deed of NASH \u2013 Nexton Holdings, LLC, dated June 11, 2025, and recorded June 13, 2025, in "
    "Book 5229, Page 36, in the Office of the Register of Deeds for Berkeley County, South Carolina "
    "(Instrument No. 2025019317).", space_after=8)

add_footer_text(doc)

# ---------- SCHEDULE B, PART I — REQUIREMENTS ----------
add_page_break(doc); add_page_header(doc)
p_center_bold(doc, "SCHEDULE B, PART I\u2014Requirements", size=12, space_before=6)
p_normal(doc, "All of the following Requirements must be met:", space_after=8)

add_numbered_item(doc, 1,
    "The Proposed Insured must notify the Company in writing of the name of any party not referred to in this "
    "Commitment who will obtain an interest in the Land or who will make a loan on the Land. The Company "
    "may then make additional Requirements or Exceptions.")

add_numbered_item(doc, 2, "Pay the agreed amount for the estate or interest to be insured.")
add_numbered_item(doc, 3, "Pay the premiums, fees, and charges for the Policy to the Company.")

add_numbered_item(doc, 4,
    "We must be furnished with a copy of SCID 3601 executed pursuant to Section 38-75-960 S.C. Code of "
    "Laws 1976, as amended, and an executed Notice of Availability of Title Insurance pursuant to S.C. "
    "Insurance Department Regulation R-69-18.")

add_numbered_item(doc, 5,
    "Seller\u2019s/Owner\u2019s Affidavit executed by Ashton Charleston Residential L.L.C. on a form to be supplied "
    "by the Company, stating that there have been no improvements to the Land within the past 90 days which "
    "could give rise to a construction lien and that there are no accounts or claims pending and unpaid which "
    "could constitute a lien against the Land.")

add_numbered_item(doc, 6, "Receipt of the acknowledged [Insurance Carrier] Privacy Policy.")

add_numbered_item(doc, 7,
    "Documents satisfactory to the Company that convey the Title or create the Mortgage to be insured, or "
    "both, must be properly authorized, executed, delivered, and recorded in the Public Records.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Inches(0.5)
r = p.add_run(
    "Deed from Ashton Charleston Residential L.L.C. to [Purchaser] conveying the Land described in "
    "Schedule A herein."
); r.font.size = Pt(10)

add_numbered_item(doc, 8,
    "Authority Documents for Ashton Charleston Residential L.L.C.: certified copy of Articles of Organization, "
    "Operating Agreement, and a current Certificate of Existence from the South Carolina Secretary of State, "
    "together with a resolution or written consent of the member(s)/manager(s) authorizing the conveyance.")

add_numbered_item(doc, 9,
    "Evidence satisfactory to the Company that the 2025 ad valorem taxes assessed under the parent parcel "
    "(TMS 195-00-00-124, Nash-Nexton Holdings, LLC) have been paid and that the subject Lot 1635 has been "
    "allocated and removed from the parent for 2026 assessment. A pro-rata tax letter or Assessor's "
    "confirmation of the new TMS 209-14-02-015 assessment status will be required at closing.")

add_numbered_item(doc, 10,
    "Roll-back taxes, if any, assessed under S.C. Code \u00a7 12-43-220 and related sections for the change in use "
    "of the subject Lot from the agricultural/undeveloped parent tract must be confirmed with the Berkeley "
    "County Assessor and, if owed, paid at closing.")

add_numbered_item(doc, 11,
    "Special assessments under the Nexton Improvement District (created by Berkeley County Ordinance "
    "No. 14-09-27) must be confirmed current; any unpaid installments or confirmed liens must be paid and "
    "released at closing.")

add_numbered_item(doc, 12,
    "Payment of the Berkeley County School District School Improvement Fee imposed under the Declaration "
    "recorded November 30, 2023, in Book 4723, Page 646 (Instrument No. 2023036406), together with a "
    "receipt evidencing payment for Lot 1635.")

add_numbered_item(doc, 13,
    "Estoppel letter or paid-assessment letter from the Nexton Residential Community Association, Inc. (or "
    "its managing agent) confirming that all regular and special assessments under the Charter for Nexton "
    "Residential Community (Book 11034, Page 153) and the First Amended and Restated Declaration of "
    "Easements and Covenant to Share Costs for Nexton (Book 11046, Page 226), as supplemented for "
    "Midtown \u2013 Phase 3C-2 (Book 5200, Page 463), are paid current as of the Commitment Date for Lot 1635.")

add_numbered_item(doc, 14,
    "Confirmation from NASH \u2013 Nexton Holdings, LLC (or its successor builder-program administrator) that "
    "the Grantor\u2019s Right of First Refusal and Right to Repurchase set forth in Exhibit C of the vesting deed "
    "(Book 5229, Page 36) has been satisfied, terminated, or waived as to Lot 1635, OR that construction of "
    "a single-family dwelling has commenced in good faith as defined therein, OR that three (3) years have "
    "elapsed since June 13, 2025. Evidence of termination to be recorded at or prior to closing.")

add_numbered_item(doc, 15,
    "Recording of all required ancillary Builder / Design-Review consents or lot-release instruments required "
    "under the Declaration of Builder Covenants (Exhibit D of the vesting deed) and the Design Guidelines "
    "issued under the Charter for Nexton Residential Community.")

add_numbered_item(doc, 16,
    "The Company may make other requirements or exceptions upon its review of the proposed documents "
    "creating the estate or interest to be insured or otherwise ascertaining details of the transaction.")

add_footer_text(doc)

# ---------- SCHEDULE B, PART II — EXCEPTIONS ----------
add_page_break(doc); add_page_header(doc)
p_center_bold(doc, "SCHEDULE B, PART II\u2014Exceptions", size=12, space_before=6)

p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    "Some historical land records contain Discriminatory Covenants that are illegal and unenforceable by law. "
    "This Commitment and the Policy treat any Discriminatory Covenant in a document referenced in Schedule "
    "B as if each Discriminatory Covenant is redacted, repudiated, removed, and not republished or recirculated. "
    "Only the remaining provisions of the document will be excepted from coverage."
); r.font.size = Pt(10); r.font.bold = True

p_normal(doc,
    "The Policy will not insure against loss or damage resulting from the terms and conditions of any lease or easement "
    "identified in Schedule A, and will include the following Exceptions unless cleared to the satisfaction of the Company:",
    space_after=8)

# Standard exceptions 1-6
add_numbered_item(doc, 1,
    "Any defect, lien, encumbrance, adverse claim, or other matter that appears for the first time in the Public "
    "Records or is created, attaches, or is disclosed between the Commitment Date and the date on which all of "
    "the Schedule B, Part I\u2014Requirements are met.")

add_numbered_item(doc, 2,
    "(a) Taxes or assessments that are not shown as existing liens by the records of any taxing authority that "
    "levies taxes or assessments on real property or by the Public Records; (b) proceedings by a public agency "
    "that may result in taxes or assessments, or notices of such proceedings, whether or not shown by the "
    "records of such agency or by the Public Records.")

add_numbered_item(doc, 3,
    "Any facts, rights, interests, or claims that are not shown by the Public Records but that could be "
    "ascertained by an inspection of the Land or that may be asserted by persons in possession of the Land.")

add_numbered_item(doc, 4, "Easements, liens or encumbrances, or claims thereof, not shown by the Public Records.")

add_numbered_item(doc, 5,
    "Any encroachment, encumbrance, violation, variation, or adverse circumstance affecting the Title that "
    "would be disclosed by an accurate and complete land survey of the Land and not shown by the Public Records.")

add_numbered_item(doc, 6, "Any mineral or mineral rights leased, granted or retained by current or prior owners.")

p_normal(doc,
    "NOTE: Exceptions 1\u20136 above will be deleted upon issuance of the Loan Policy Only.",
    bold=True, space_before=6, space_after=10)

# Taxes
add_numbered_item(doc, 7,
    "Taxes and assessments for the year 2026, and subsequent years, not yet due and payable. 2025 ad valorem "
    "taxes for the subject Lot remain assessed under the parent parcel TMS 195-00-00-124 (Nash-Nexton "
    "Holdings, LLC \u2013 1,105.13 acres at 1283 State Road); parent-parcel 2025 taxes were paid October 24, 2025. "
    "Separate assessment for TMS 209-14-02-015 (Lot 1635) has not yet been issued by the Berkeley County "
    "Assessor. Subject to roll-back taxes, if applicable, under S.C. Code \u00a7 12-43-220.")

# Nexton Improvement District
add_numbered_item(doc, 8,
    "Special assessments under the Nexton Improvement District authorized and created by Berkeley County "
    "Ordinance No. 14-09-27, together with any supplements, amendments, and levies issued thereunder.")

# Governing plats
add_group_heading(doc, "Governing Subdivision Plats")

add_numbered_item(doc, 9,
    "Those certain matters, easements, rights-of-way, setbacks, HOA areas, drainage areas, and utility "
    "corridors shown and described on the Subdivision Plat of Midtown \u2013 Phase 3C-2 (85 lots) of Tract 2A, "
    "TMS 195-00-00-124, prepared by Johnathan F. Burns, PLS #22742, of GPA, Inc., dated November 14, "
    "2024, and recorded April 3, 2025, at Instrument Nos. 2025010343\u20132025010346 in the Office of the "
    "Register of Deeds for Berkeley County, South Carolina (including Ivy Green Lane, June Berry Drive, "
    "Twinflower Lane, Wood Lilly Loop, HOA Areas 81\u201383, Private R/W No. 74, and new BCWS, BEC, and "
    "drainage easements).")

# ---- RESTRICTIONS & AGREEMENTS (grouped by common association, oldest → newest) ----
p_center_bold(doc, "Restrictions, Covenants, and Development Agreements", size=10, space_before=12, space_after=4)

# Group: Nexton Master Development / HOA Regime
add_group_heading(doc, "Nexton Master Development Agreement and HOA Regime (oldest to most recent)")

add_numbered_item(doc, 10,
    "Nexton Development Agreement (formerly known as The Parks of Berkeley Development Agreement) "
    "by and between Berkeley County, South Carolina and MWV-Parks of Berkeley, LLC, dated April 3, "
    "2006, and recorded April 21, 2006, in Book 5549, Page 1, in the Office of the Register of Deeds for "
    "Berkeley County, South Carolina, together with any Addendum, Amendments, Second Amendment, "
    "and all subsequent amendments, supplements, and assignments thereto.")

add_numbered_item(doc, 11,
    "Declaration of Easements and Covenant to Share Costs for Nexton, dated and recorded January 18, "
    "2013, in Book 9906, Page 285 (Instrument No. 2013-00001654), in the Office of the Register of Deeds "
    "for Berkeley County, South Carolina.")

add_numbered_item(doc, 12,
    "Charter for Nexton Residential Community, dated October 22, 2014, and recorded October 24, 2014, "
    "in Book 11034, Page 153 (Instrument No. 2014-00024763), in the Office of the Register of Deeds for "
    "Berkeley County, South Carolina, together with all amendments and supplements thereto.")

add_numbered_item(doc, 13,
    "First Amended and Restated Declaration of Easements and Covenant to Share Costs for Nexton, dated "
    "and recorded October 31, 2014, in Book 11046, Page 226, in the Office of the Register of Deeds for "
    "Berkeley County, South Carolina, together with all amendments and supplements thereto.")

add_numbered_item(doc, 14,
    "Supplement to the Charter for Nexton Residential Community and the Covenant to Share Costs "
    "(Midtown \u2013 Phase 3C-2) by NASH \u2013 Nexton Holdings, LLC, recorded May 19, 2025, in Book 5200, "
    "Page 463 (Instrument No. 2025015917), in the Office of the Register of Deeds for Berkeley County, "
    "South Carolina.")

add_numbered_item(doc, 15,
    "Partial Assignment and Assumption of Rights and Obligations Under Development Agreement between "
    "NASH \u2013 Nexton Holdings, LLC and Ashton Charleston Residential L.L.C., dated June 13, 2025, and "
    "recorded June 13, 2025, in Book 5229, Page 59 (Instrument No. 2025019318), in the Office of the "
    "Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 16,
    "Agreement of Covenants, Conditions and Restrictions (including Grantor\u2019s Right of First Refusal and "
    "Grantor\u2019s Right to Repurchase) set forth as Exhibit C of the vesting Limited Warranty Deed recorded in "
    "Book 5229, Page 36, and the Declaration of Builder Covenants set forth as Exhibit D thereto.")

# Group: Wetlands
add_group_heading(doc, "Wetlands Restrictive Covenants (oldest to most recent)")

add_numbered_item(doc, 17,
    "Declaration of Restrictive Covenants (Wetlands), recorded December 16, 2005, in Book 5235, Page 75, "
    "in the Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 18,
    "Declaration of Restrictive Covenants (Wetlands), recorded November 10, 2011, in Book 9183, Page "
    "160, in the Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 19,
    "Declaration of Restrictive Covenants (Wetlands) by MWV-Parks of Berkeley, LLC, dated April 2, "
    "2013, and recorded June 25, 2013, in Book 10205, Page 1, in the Office of the Register of Deeds for "
    "Berkeley County, South Carolina.")

# Group: Stormwater Maintenance
add_group_heading(doc, "Permanent Stormwater Maintenance Covenants (oldest to most recent)")

add_numbered_item(doc, 20,
    "Covenants for Permanent Maintenance of Stormwater Systems, recorded March 8, 2023, in Book 4500, "
    "Page 143, in the Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 21,
    "Covenants for Permanent Maintenance of Stormwater Systems, recorded January 4, 2024, in Book "
    "4749, Page 719, in the Office of the Register of Deeds for Berkeley County, South Carolina.")

# Group: Berkeley County / School / Utility Agreements
add_group_heading(doc, "Berkeley County, School District, and Santee Cooper Agreements (oldest to most recent)")

add_numbered_item(doc, 22,
    "Berkeley County Resolution regarding the Nexton development, recorded May 18, 2023, in Book 4559, "
    "Page 433, in the Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 23,
    "Declaration of School Improvement Fee between Berkeley County School District and NASH \u2013 Nexton "
    "Holdings, LLC, dated November 29, 2023, and recorded November 30, 2023, in Book 4723, Page 646 "
    "(Instrument No. 2023036406), in the Office of the Register of Deeds for Berkeley County, South "
    "Carolina, imposing a per-lot School Improvement Fee payable in connection with development of "
    "each residential lot.")

add_numbered_item(doc, 24,
    "Agreement between NASH \u2013 Nexton Holdings, LLC and the South Carolina Public Service Authority "
    "(Santee Cooper), recorded January 17, 2024, in Book 4756, Page 760 (Instrument No. 2024001397), in "
    "the Office of the Register of Deeds for Berkeley County, South Carolina.")

# Group: Telecommunications Master Agreement
add_group_heading(doc, "Telecommunications Master Agreement")

add_numbered_item(doc, 25,
    "Memorandum of Master Agreement by and between MWV-Parks of Berkeley, LLC, NextIP, LLC and "
    "Berkeley Cable Television Company, Inc., dated December 18, 2013, as evidenced by Memorandum "
    "dated September 21, 2015, and recorded September 29, 2015, in Book 2026, Page 411 (Instrument No. "
    "2015033183), in the Office of the Register of Deeds for Berkeley County, South Carolina, together with "
    "all amendments, supplements, and assignments thereto.")

# ---- EASEMENTS (grouped by common association, oldest → newest) ----
p_center_bold(doc, "Easements and Rights-of-Way", size=10, space_before=12, space_after=4)

# Group: Pre-Nexton (Westvaco / Sheep Island parent tract)
add_group_heading(doc, "Pre-Nexton Utility Easements \u2013 Former Westvaco / Sheep Island Parent Tract (oldest to most recent)")

add_numbered_item(doc, 26,
    "Grant of Perpetual Easement to the South Carolina Public Service Authority (Santee Cooper), recorded "
    "in Book 234, Page 180, in the Office of the Register of Deeds for Berkeley County, South Carolina, "
    "being a drainage easement against the former parent tract.")

add_numbered_item(doc, 27,
    "Grant of Perpetual Easement to the South Carolina Public Service Authority (Santee Cooper) for an "
    "underground water transmission system (Easement No. 177, Drawing No. 8001-B03-5110), recorded in "
    "Book 343, Page 280, in the Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 28,
    "Power Line Easement from Westvaco Corporation (formerly West Virginia Pulp & Paper Company) to "
    "the South Carolina Public Service Authority (Santee Cooper), recorded in Book C150, Page 204, in the "
    "Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 29,
    "Power Line Easement from Westvaco Corporation to the South Carolina Public Service Authority "
    "(Santee Cooper), dated and recorded January 10, 1983, in Book C152, Page 60, in the Office of the "
    "Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 30,
    "Power Line Easement from Westvaco Corporation to Central Electric Power Cooperative, Inc., recorded "
    "June 15, 1984, in Book C167, Page 278, in the Office of the Register of Deeds for Berkeley County, "
    "South Carolina.")

add_numbered_item(doc, 31,
    "Right-of-Way Easement from Westvaco Corporation to Summerville Commissioners of Public Works, "
    "recorded July 8, 1992, in Book 124, Page 312, in the Office of the Register of Deeds for Berkeley "
    "County, South Carolina, for a 15-foot-wide water line crossing the former Sheep Island Main Tract.")

# Group: Nexton-era electric / telecom
add_group_heading(doc, "Nexton-Era Electric and Telecommunication Easements (oldest to most recent)")

add_numbered_item(doc, 32,
    "Right of Way Easement from MWV-Sheep Island, LLC to Home Telephone ILEC, LLC, d/b/a Home "
    "Telecom, dated March 18, 2011, and recorded March 22, 2011, in Book 8865, Page 110, in the Office "
    "of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 33,
    "Agreement between MWV-Sheep Island, LLC and Berkeley Electric Cooperative, Inc., dated December "
    "29, 2010, and recorded May 31, 2011, in Book 8954, Page 45, in the Office of the Register of Deeds "
    "for Berkeley County, South Carolina.")

add_numbered_item(doc, 34,
    "Right-of-Way Easement to South Carolina Electric & Gas Company (SCE&G, now Dominion Energy "
    "South Carolina), recorded November 17, 2011, in Book 9192, Page 189, in the Office of the Register of "
    "Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 35,
    "Right-of-Way Easement to South Carolina Electric & Gas Company (SCE&G, now Dominion Energy "
    "South Carolina), recorded January 31, 2012, in Book 9297, Page 40, in the Office of the Register of "
    "Deeds for Berkeley County, South Carolina.")

# Group: Berkeley County Water & Sanitation (Phase 3C utility)
add_group_heading(doc, "Berkeley County Water & Sanitation \u2013 Phase 3C Utility Easements and Conveyances (oldest to most recent)")

add_numbered_item(doc, 36,
    "Grant of Perpetual Exclusive Easement to Berkeley County (in care of Berkeley County Water & "
    "Sanitation) from NASH \u2013 Nexton Holdings, LLC, recorded May 12, 2025, in Book 5193, Page 787 "
    "(Instrument No. 2025015075), in the Office of the Register of Deeds for Berkeley County, South "
    "Carolina, for the \u201cNEW 2.5\u2019 BCWS EXCL. U.E.\u201d, \u201cNEW 5\u2019 BCWS EXCL. U.E.\u201d, and \u201cNEW VAR. "
    "WIDTH BCWS U.E.\u201d corridors as shown on the governing Phase 3C plats.")

add_numbered_item(doc, 37,
    "Grant of Perpetual Non-Exclusive Easement to Berkeley County (in care of Berkeley County Water & "
    "Sanitation) from NASH \u2013 Nexton Holdings, LLC, recorded May 12, 2025, in Book 5193, Page 793 "
    "(Instrument No. 2025015076), in the Office of the Register of Deeds for Berkeley County, South "
    "Carolina, for underground water and sewer lines across the strips shown on the governing Phase 3C plats.")

add_numbered_item(doc, 38,
    "Bill of Sale from NASH \u2013 Nexton Holdings, LLC to Berkeley County Water & Sanitation, recorded "
    "October 9, 2025, in Book 5348, Page 6 (Instrument No. 2025034058), in the Office of the Register of "
    "Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 39,
    "Title to Sewer System from NASH \u2013 Nexton Holdings, LLC to Berkeley County Water & Sanitation for "
    "Nexton Midtown Phase 3C-2, recorded October 9, 2025, in Book 5348, Page 10 (Instrument No. "
    "2025034059), in the Office of the Register of Deeds for Berkeley County, South Carolina.")

add_numbered_item(doc, 40,
    "Utility Deed from NASH \u2013 Nexton Holdings, LLC to Berkeley County Water & Sanitation, recorded "
    "October 9, 2025, in Book 5348, Page 14 (Instrument No. 2025034060), in the Office of the Register of "
    "Deeds for Berkeley County, South Carolina.")

# Catch-all
add_numbered_item(doc, 41,
    "Rights of upper and lower riparian owners in and to the waters of streams, creeks, or branches crossing "
    "or adjoining the Land, if any, and the natural flow thereof free from diminution or pollution.")

add_numbered_item(doc, 42,
    "Zoning, subdivision, and other land-use laws, regulations, and ordinances applicable to the Land.")

add_numbered_item(doc, 43,
    "Such matters as would be disclosed by an accurate survey and inspection of the Land.")

add_numbered_item(doc, 44,
    "The Company may make other requirements or exceptions upon its review of the proposed documents "
    "creating the estate or interest to be insured or otherwise ascertaining details of the transaction.")

add_footer_text(doc)

# ---------- SAVE ----------
doc.save(OUTPUT)
print(f"Done! Saved to: {OUTPUT}")
