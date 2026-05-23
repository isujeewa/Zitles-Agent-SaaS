#!/usr/bin/env python3
"""Generate a draft ALTA Title Commitment (2021 form) for 145 Oyster Point Row."""

from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "commitment-4260400003.docx")


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 6, "val": "single", "color": "000000"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            element = OxmlElement(f"w:{edge}")
            for attr, val in kwargs[edge].items():
                element.set(qn(f"w:{attr}"), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_header_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_right_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(9)


def add_page_header(doc):
    """Standard page header — right-aligned commitment info."""
    add_right_header(doc, "Commitment for Title Insurance")
    add_right_header(doc, "South Carolina - 2021 v. 01.00 (07-01-2021)")
    add_header_line(doc)


def add_footer_text(doc):
    """Add the standard ALTA footer disclaimer as small italic text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "999999")
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(
        "This page is only a part of a 2021 ALTA Commitment for Title Insurance. "
        "This Commitment is not valid without the Notice; the Commitment to Issue Policy; "
        "the Commitment Conditions; Schedule A; Schedule B, Part I\u2014Requirements; "
        "Schedule B, Part II\u2014Exceptions; and a counter-signature by the Company or "
        "its issuing agent that may be in electronic form."
    )
    run.font.size = Pt(7)
    run.font.italic = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    pPr2 = p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2 = OxmlElement("w:top")
    top2.set(qn("w:val"), "single")
    top2.set(qn("w:sz"), "4")
    top2.set(qn("w:space"), "1")
    top2.set(qn("w:color"), "000000")
    pBdr2.append(top2)
    pPr2.append(pBdr2)
    run2 = p2.add_run(
        "Copyright 2021 American Land Title Association. All rights reserved.\n"
        "The use of this Form (or any derivative thereof) is restricted to ALTA licensees and\n"
        "ALTA members in good standing as of the date of use. All other uses are prohibited.\n"
        "Reprinted under license from the American Land Title Association."
    )
    run2.font.size = Pt(7)
    run2.font.bold = True


def p_normal(doc, text, bold=False, size=10, space_after=6, space_before=0, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def p_center_bold(doc, text, size=12, space_before=12, space_after=8):
    return p_normal(doc, text, bold=True, size=size, space_after=space_after,
                    space_before=space_before, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_numbered_item(doc, number, text, bold_number=True, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    run_num = p.add_run(f"{number}.\t")
    run_num.font.size = Pt(10)
    run_num.font.bold = bold_number
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    return p


def add_sub_item(doc, letter, text, indent=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run_letter = p.add_run(f"{letter}.\t")
    run_letter.font.size = Pt(10)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    return p


# ============================================================
#  BUILD THE DOCUMENT
# ============================================================
doc = Document()
set_narrow_margins(doc)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)

# ============================================================
#  PAGE 1 — NOTICE / COMMITMENT TO ISSUE POLICY
# ============================================================
add_page_header(doc)

p_center_bold(doc, "ALTA COMMITMENT FOR TITLE INSURANCE\nissued by\n[INSURANCE CARRIER]", size=11, space_before=24)

p_center_bold(doc, "NOTICE", size=11, space_before=18)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("IMPORTANT\u2014READ CAREFULLY")
run.font.size = Pt(10)
run.font.bold = True
run2 = p.add_run(
    ": THIS COMMITMENT IS AN OFFER TO ISSUE ONE OR MORE TITLE INSURANCE POLICIES. "
    "ALL CLAIMS OR REMEDIES SOUGHT AGAINST THE COMPANY INVOLVING THE CONTENT OF THIS "
    "COMMITMENT OR THE POLICY MUST BE BASED SOLELY IN CONTRACT."
)
run2.font.size = Pt(10)

p_normal(doc,
    "THIS COMMITMENT IS NOT AN ABSTRACT OF TITLE, REPORT OF THE CONDITION OF TITLE, LEGAL "
    "OPINION, OPINION OF TITLE, OR OTHER REPRESENTATION OF THE STATUS OF TITLE. THE "
    "PROCEDURES USED BY THE COMPANY TO DETERMINE INSURABILITY OF THE TITLE, INCLUDING ANY "
    "SEARCH AND EXAMINATION, ARE PROPRIETARY TO THE COMPANY, WERE PERFORMED SOLELY FOR "
    "THE BENEFIT OF THE COMPANY, AND CREATE NO EXTRACONTRACTUAL LIABILITY TO ANY PERSON, "
    "INCLUDING A PROPOSED INSURED.", size=10, space_after=8)

p_normal(doc,
    "THE COMPANY\u2019S OBLIGATION UNDER THIS COMMITMENT IS TO ISSUE A POLICY TO A PROPOSED "
    "INSURED IDENTIFIED IN SCHEDULE A IN ACCORDANCE WITH THE TERMS AND CONDITIONS OF THIS "
    "COMMITMENT. THE COMPANY HAS NO LIABILITY OR OBLIGATION INVOLVING THE CONTENT OF THIS "
    "COMMITMENT TO ANY OTHER PERSON.", size=10, space_after=12)

p_center_bold(doc, "COMMITMENT TO ISSUE POLICY", size=11)

p_normal(doc,
    'Subject to the Notice; Schedule B, Part I\u2014Requirements; Schedule B, Part II\u2014Exceptions; '
    'and the Commitment Conditions, [Insurance Carrier], a [State] Corporation (the \u201cCompany\u201d), '
    'commits to issue the Policy according to the terms and provisions of this Commitment. This '
    'Commitment is effective as of the Commitment Date shown in Schedule A for each Policy described '
    'in Schedule A, only when the Company has entered in Schedule A both the specified dollar amount '
    'as the Proposed Amount of Insurance and the name of the Proposed Insured.', space_after=8)

p_normal(doc,
    "If all of the Schedule B, Part I\u2014Requirements have not been met within 180 days after the "
    "Commitment Date, this Commitment terminates and the Company\u2019s liability and obligation end.",
    space_after=12)

p_normal(doc, "[INSURANCE CARRIER]", bold=True, space_after=18)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("By: ____________________________\t\t\tBy: ____________________________")
run.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("      [Name], President\t\t\t\t\t      [Name], Secretary")
run2.font.size = Pt(9)
run2.font.bold = True

add_footer_text(doc)

# ============================================================
#  PAGE 2-4 — COMMITMENT CONDITIONS
# ============================================================
add_page_break(doc)
add_page_header(doc)
p_center_bold(doc, "COMMITMENT CONDITIONS", size=11, space_before=6)

add_numbered_item(doc, 1, "DEFINITIONS")
definitions = [
    ("a", '"Discriminatory Covenant": Any covenant, condition, restriction, or limitation that is unenforceable under applicable law because it illegally discriminates against a class of individuals based on personal characteristics such as race, color, religion, sex, sexual orientation, gender identity, familial status, disability, national origin, or other legally protected class.'),
    ("b", '"Knowledge" or "Known": Actual knowledge or actual notice, but not constructive notice imparted by the Public Records.'),
    ("c", '"Land": The land described in Item 5 of Schedule A and improvements located on that land that by State law constitute real property. The term "Land" does not include any property beyond that described in Schedule A, nor any right, title, interest, estate, or easement in any abutting street, road, avenue, alley, lane, right-of-way, body of water, or waterway, but does not modify or limit a right of access to and from the Land is to be insured by the Policy.'),
    ("d", '"Mortgage": A mortgage, deed of trust, trust deed, security deed, or other real property security instrument, including one evidenced by electronic means authorized by law.'),
    ("e", '"Policy": Each contract of title insurance, in a form adopted by the American Land Title Association, issued or to be issued by the Company pursuant to this Commitment.'),
    ("f", '"Proposed Amount of Insurance": Each dollar amount specified in Schedule A as the Proposed Amount of Insurance of each Policy to be issued pursuant to this Commitment.'),
    ("g", '"Proposed Insured": Each person identified in Schedule A as the Proposed Insured of each Policy to be issued pursuant to this Commitment.'),
    ("h", '"Public Records": The recording or filing system established under State statutes in effect at the Commitment Date under which a document must be recorded or filed to impart constructive notice of matters relating to the Title to a purchaser for value without Knowledge. The term "Public Records" does not include any other recording or filing system, including any pertaining to environmental remediation or protection, planning, permitting, zoning, licensing, building, health, public safety, or national security matters.'),
    ("i", '"State": The state or commonwealth of the United States within whose exterior boundaries the Land is located. The term "State" also includes the District of Columbia, the Commonwealth of Puerto Rico, the U.S. Virgin Islands, and Guam.'),
    ("j", '"Title": The estate or interest in the Land identified in Item 3 of Schedule A.'),
]
for letter, text in definitions:
    add_sub_item(doc, letter, text)

add_numbered_item(doc, 2,
    "If all of the Schedule B, Part I\u2014Requirements have not been met within the time period specified "
    "in the Commitment to Issue Policy, this Commitment terminates and the Company\u2019s liability and obligation end.")

add_numbered_item(doc, 3, "The Company\u2019s liability and obligation is limited by and this Commitment is not valid without:")
for letter, text in [("a", "the Notice;"), ("b", "the Commitment to Issue Policy;"),
                      ("c", "the Commitment Conditions;"), ("d", "Schedule A;"),
                      ("e", "Schedule B, Part I\u2014Requirements;"),
                      ("f", "Schedule B, Part II\u2014Exceptions; and"),
                      ("g", "a counter-signature by the Company or its issuing agent that may be in electronic form.")]:
    add_sub_item(doc, letter, text)

add_page_break(doc)
add_page_header(doc)

add_numbered_item(doc, 4, "COMPANY\u2019S RIGHT TO AMEND")
p_normal(doc,
    "The Company may amend this Commitment at any time. If the Company amends this Commitment to add "
    "a defect, lien, encumbrance, adverse claim, or other matter recorded in the Public Records prior to the "
    "Commitment Date, any liability of the Company is limited by Commitment Condition 5. The Company is not "
    "liable for any other amendment to this Commitment.", space_after=8)

add_numbered_item(doc, 5, "LIMITATIONS OF LIABILITY")
add_sub_item(doc, "a",
    "The Company\u2019s liability under Commitment Condition 4 is limited to the Proposed Insured\u2019s actual "
    "expense incurred in the interval between the Company\u2019s delivery to the Proposed Insured of the "
    "Commitment and the delivery of the amended Commitment, resulting from the Proposed Insured\u2019s good "
    "faith reliance to:")
for sub in ["comply with the Schedule B, Part I\u2014Requirements;",
            "eliminate, with the Company\u2019s written consent, any Schedule B, Part II\u2014Exceptions; or",
            "acquire the Title or create the Mortgage covered by this Commitment."]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(1.5)
    run = p.add_run(sub)
    run.font.size = Pt(10)

add_sub_item(doc, "b",
    "The Company is not liable under Commitment Condition 5.a. if the Proposed Insured requested the "
    "amendment or had Knowledge of the matter and did not notify the Company about it in writing.")
add_sub_item(doc, "c",
    "The Company is only liable under Commitment Condition 4 if the Proposed Insured would not have "
    "incurred the expense had the Commitment included the added matter when the Commitment was first "
    "delivered to the Proposed Insured.")
add_sub_item(doc, "d",
    "The Company\u2019s liability does not exceed the lesser of the Proposed Insured\u2019s actual expense incurred "
    "in good faith and described in Commitment Condition 5.a. or the Proposed Amount of Insurance.")
add_sub_item(doc, "e",
    "The Company is not liable for the content of the Transaction Identification Data, if any.")
add_sub_item(doc, "f",
    "The Company is not obligated to issue the Policy referred to in this Commitment unless all of the "
    "Schedule B, Part I\u2014Requirements have been met to the satisfaction of the Company.")
add_sub_item(doc, "g",
    "The Company\u2019s liability is further limited by the terms and provisions of the Policy to be issued to the "
    "Proposed Insured.")

add_numbered_item(doc, 6,
    "LIABILITY OF THE COMPANY MUST BE BASED ON THIS COMMITMENT; CHOICE OF LAW AND CHOICE OF FORUM")
add_sub_item(doc, "a",
    "Only a Proposed Insured identified in Schedule A, and no other person, may make a claim under this Commitment.")
add_sub_item(doc, "b",
    "Any claim must be based in contract under the State law of the State where the Land is located and is "
    "restricted to the terms and provisions of this Commitment. Any litigation or other proceeding brought "
    "by the Proposed Insured against the Company must be filed only in a State or federal court having jurisdiction.")
add_sub_item(doc, "c",
    "This Commitment, as last revised, is the exclusive and entire agreement between the parties with "
    "respect to the subject matter of this Commitment and supersedes all prior commitment negotiations, "
    "representations, and proposals of any kind, whether written or oral, express or implied, relating to the "
    "subject matter of this Commitment.")
add_sub_item(doc, "d",
    "The deletion or modification of any Schedule B, Part II\u2014Exception does not constitute an agreement "
    "or obligation to provide coverage beyond the terms and provisions of this Commitment or the Policy.")
add_sub_item(doc, "e",
    "Any amendment or endorsement to this Commitment must be in writing and authenticated by a person "
    "authorized by the Company.")
add_sub_item(doc, "f",
    "When the Policy is issued, all liability and obligation under this Commitment will end and the Company\u2019s "
    "only liability will be under the Policy.")

add_page_break(doc)
add_page_header(doc)

add_numbered_item(doc, 7, "IF THIS COMMITMENT IS ISSUED BY AN ISSUING AGENT")
p_normal(doc,
    "The issuing agent is the Company\u2019s agent only for the limited purpose of issuing title insurance commitments "
    "and policies. The issuing agent is not the Company\u2019s agent for closing, settlement, escrow, or any other purpose.",
    space_after=8)

add_numbered_item(doc, 8, "PRO-FORMA POLICY")
p_normal(doc,
    "The Company may provide, at the request of a Proposed Insured, a pro-forma policy illustrating the coverage "
    "that the Company may provide. A pro-forma policy neither reflects the status of Title at the time that the pro-forma "
    "policy is delivered to a Proposed Insured, nor is it a commitment to insure.", space_after=8)

add_numbered_item(doc, 9, "CLAIMS PROCEDURES")
p_normal(doc,
    "This Commitment incorporates by reference all Conditions for making a claim in the Policy to be issued to the "
    "Proposed Insured. Commitment Condition 9 does not modify the limitations of liability in Commitment Conditions 5 and 6.",
    space_after=8)

add_numbered_item(doc, 10, "CLASS ACTION")
p_normal(doc,
    "ALL CLAIMS AND DISPUTES ARISING OUT OF OR RELATING TO THIS COMMITMENT, INCLUDING ANY SERVICE OR OTHER "
    "MATTER IN CONNECTION WITH ISSUING THIS COMMITMENT, ANY BREACH OF A COMMITMENT PROVISION, OR ANY OTHER "
    "CLAIM OR DISPUTE ARISING OUT OF OR RELATING TO THE TRANSACTION GIVING RISE TO THIS COMMITMENT, MUST BE "
    "BROUGHT IN AN INDIVIDUAL CAPACITY. NO PARTY MAY SERVE AS PLAINTIFF, CLASS MEMBER, OR PARTICIPANT IN ANY "
    "CLASS OR REPRESENTATIVE PROCEEDING. ANY POLICY ISSUED PURSUANT TO THIS COMMITMENT WILL CONTAIN A CLASS "
    "ACTION CONDITION.", size=10, space_after=8)

add_numbered_item(doc, 11, "ARBITRATION")
p_normal(doc,
    "The Policy contains an arbitration clause. All arbitrable matters when the Proposed Amount of Insurance is "
    "$2,000,000 or less may be arbitrated at the election of either the Company or the Proposed Insured as the "
    "exclusive remedy of the parties. A Proposed Insured may review a copy of the arbitration rules at "
    "http://www.alta.org/arbitration.", space_after=8)

add_footer_text(doc)

# ============================================================
#  PAGE 5 — TRANSACTION ID + SCHEDULE A
# ============================================================
add_page_break(doc)
add_page_header(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Transaction Identification Data, for which the Company assumes no liability as set forth in Commitment Condition 5.e.:")
run.font.size = Pt(10)
run.font.bold = True

tid_lines = [
    ("Issuing Agent:", "[Title Agency Name]"),
    ("Issuing Office:", "[Address]"),
    ("", "[City, State ZIP]"),
    ("Issuing Office\u2019s ALTA\u00ae Registry ID:", ""),
    ("Loan ID Number:", ""),
    ("Commitment Number:", "[File Number]"),
    ("Issuing Office File Number:", "[File Number]"),
    ("Property Address:", "145 Oyster Point Row, Charleston, SC 29412"),
    ("Revision Number:", ""),
]
for label, value in tid_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if label:
        run_l = p.add_run(f"{label} ")
        run_l.font.size = Pt(10)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)

p_center_bold(doc, "SCHEDULE A", size=12, space_before=18)

add_numbered_item(doc, 1, "Commitment Date:  at 8:00 AM")

add_numbered_item(doc, 2, "Policy to be issued:")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = Inches(1.0)
lines = [
    ("a.\t2021 ALTA Homeowner\u2019s Policy", ""),
    ("\tProposed Insured:", "\t\t[Proposed Insured]"),
    ("\tProposed Amount of Insurance:", "\t$___________"),
    ("\tThe estate or interest to be insured:", "\tfee simple"),
]
for left, right in lines:
    run = p.add_run(f"{left}{right}\n")
    run.font.size = Pt(10)

add_numbered_item(doc, 3, "The estate or interest in the Land at the Commitment Date is:\n\tfee simple")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.first_line_indent = Inches(-0.35)
run_num = p.add_run("4.\t")
run_num.font.size = Pt(10)
run_num.font.bold = True
run_text = p.add_run(
    "The Title is, at the Commitment Date, vested in:\n"
    "\tKeith T. Borg and Kathy H. Borg by deed from Kerri A. Kolehma and J. Rebecca McSwain "
    "dated 07/06/2006 and recorded with Charleston County Register of Deeds on 07/07/2006 "
    "in Book E590, Page 103."
)
run_text.font.size = Pt(10)

add_numbered_item(doc, 5, "The Land is described as follows:")
p_normal(doc, "SEE EXHIBIT A ATTACHED HERETO AND MADE A PART HEREOF",
         bold=True, space_before=6, space_after=18)

# Title Agency block
p_normal(doc, "[TITLE AGENCY NAME]", bold=True, space_after=2)
p_normal(doc, "[Address]\n[City, State ZIP]\nTelephone:", space_after=12)
p_normal(doc, "Countersigned:", space_after=12)

p = doc.add_paragraph()
run = p.add_run("By:____________________________")
run.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("      Authorized Signatory")
run2.font.size = Pt(10)
run2.font.bold = True

p3 = doc.add_paragraph()
run3 = p3.add_run("[Name], License #[______]\n[Title Agency], License #")
run3.font.size = Pt(10)

add_footer_text(doc)

# ============================================================
#  PAGE 6 — EXHIBIT A
# ============================================================
add_page_break(doc)
add_page_header(doc)

p_center_bold(doc, "EXHIBIT A", size=12, space_before=6)

p_normal(doc, "Commitment No.: [File Number]", bold=True, space_after=12)

p_normal(doc,
    "The land referred to herein below is situated in the County of Charleston, State of South Carolina, "
    "and is described as follows:", space_after=12)

p_normal(doc,
    "Lot No. 71, Block B, Rivers Point Subdivision, as shown on a plat made by C. Roger Jennings, "
    "Surveyor, dated May, 1980, and recorded at the Charleston County R.M.C. Office in Plat Book AQ, "
    "at Page 31.",
    space_after=8)

p_normal(doc,
    "City of Charleston, County of Charleston, State of South Carolina.",
    space_after=8)

p_normal(doc, "TMS: 426-04-00-003", space_after=8)

p_normal(doc,
    "This being the same property conveyed to Keith T. Borg and Kathy H. Borg by deed of Kerri A. "
    "Kolehma and J. Rebecca McSwain, dated July 6, 2006, and recorded at the Charleston County R.M.C. "
    "Office on July 7, 2006, in Book E590, at Page 103.",
    space_after=8)

add_footer_text(doc)

# ============================================================
#  PAGE 7 — SCHEDULE B, PART I — REQUIREMENTS
# ============================================================
add_page_break(doc)
add_page_header(doc)

p_center_bold(doc, "SCHEDULE B, PART I\u2014Requirements", size=12, space_before=6)

p_normal(doc, "All of the following Requirements must be met:", space_after=8)

add_numbered_item(doc, 1,
    "The Proposed Insured must notify the Company in writing of the name of any party not referred to in this "
    "Commitment who will obtain an interest in the Land or who will make a loan on the Land. The Company "
    "may then make additional Requirements or Exceptions.")

add_numbered_item(doc, 2,
    "Pay the agreed amount for the estate or interest to be insured.")

add_numbered_item(doc, 3,
    "Pay the premiums, fees, and charges for the Policy to the Company.")

add_numbered_item(doc, 4,
    "We must be furnished with a copy of SCID 3601 executed pursuant to Section 38-75-960 S.C. Code of "
    "Laws 1976, as amended, and an executed Notice of Availability of Title Insurance pursuant to S.C. "
    "Insurance Department Regulation R-69-18, Vol. 25A of S.C. Code of Laws 1976, as amended.")

add_numbered_item(doc, 5,
    "Seller\u2019s/Owner\u2019s Affidavit Indemnity executed by current owner(s) of the land on a form to be supplied by "
    "the Company stating that there have been no improvements to the land within the past 90 days which "
    "could give rise to a construction lien and that there are no accounts or claims pending and unpaid which "
    "could constitute a lien against the land. The affidavit will also state that affiant has no knowledge of any "
    "natural person or legal entity who has or could have a claim of right, interest or lien adverse to the Insured.")

add_numbered_item(doc, 6,
    "Receipt of the acknowledged [Insurance Carrier] Privacy Policy.")

add_numbered_item(doc, 7,
    "Documents satisfactory to the Company that convey the Title or create the Mortgage to be insured, or "
    "both, must be properly authorized, executed, delivered, and recorded in the Public Records.")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run(
    "Warranty Deed from Keith T. Borg and Kathy H. Borg to [Purchaser] conveying the land described "
    "in Schedule A herein.")
run.font.size = Pt(10)

add_numbered_item(doc, 8,
    "Payoff, satisfaction and release of the mortgage recorded in Book 0487, Page 477 from Keith T. Borg "
    "and Kathy H. Borg to SunTrust Mortgage, Inc. (MERS as nominee), recorded on June 25, 2015 in the "
    "Office of the Charleston County Register of Deeds.")

add_numbered_item(doc, 9,
    "Satisfaction or release of the SC Department of Revenue State Tax Lien #1479166 against Borg, Keith, "
    "filed December 19, 2025, in the amount of $390.10 (Individual Income Tax, period ending 12/31/2024).")

add_numbered_item(doc, 10,
    "Quit Claim Deed or other instrument from Keith Thomas Borg to Kathy Hendrickson Borg (or current "
    "grantee) as required by the Decree of Divorce entered August 30, 2017, Case No. 2017-DR-10-2419, "
    "Family Court of the Ninth Judicial Circuit, Charleston County, South Carolina, wherein Wife was "
    "awarded sole ownership of the marital home at 145 Oyster Point Row. No such deed has been located "
    "of record.")

add_footer_text(doc)

# ============================================================
#  PAGE 8-9 — SCHEDULE B, PART II — EXCEPTIONS
# ============================================================
add_page_break(doc)
add_page_header(doc)

p_center_bold(doc, "SCHEDULE B, PART II\u2014Exceptions", size=12, space_before=6)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "Some historical land records contain Discriminatory Covenants that are illegal and unenforceable by law. "
    "This Commitment and the Policy treat any Discriminatory Covenant in a document referenced in Schedule "
    "B as if each Discriminatory Covenant is redacted, repudiated, removed, and not republished or recirculated. "
    "Only the remaining provisions of the document will be excepted from coverage."
)
run.font.size = Pt(10)
run.font.bold = True

p_normal(doc,
    "The Policy will not insure against loss or damage resulting from the terms and conditions of any lease or easement "
    "identified in Schedule A, and will include the following Exceptions unless cleared to the satisfaction of the Company:",
    space_after=8)

# Standard exceptions 1-6
add_numbered_item(doc, 1,
    "Any defect, lien, encumbrance, adverse claim, or other matter that appears for the first time in the Public "
    "Records or is created, attaches, or is disclosed between the Commitment Date and the date on which all of "
    "the Schedule B, Part I\u2014Requirements are met.")

add_numbered_item(doc, 2,
    "(a) Taxes or assessments that are not shown as existing liens by the records of any taxing authority that "
    "levies taxes or assessments on real property or by the Public Records; (b) proceedings by a public agency "
    "that may result in taxes or assessments, or notices of such proceedings, whether or not shown by the "
    "records of such agency or by the Public Records.")

add_numbered_item(doc, 3,
    "Any facts, rights, interests, or claims that are not shown by the Public Records but that could be "
    "ascertained by an inspection of the Land or that may be asserted by persons in possession in the Land.")

add_numbered_item(doc, 4,
    "Easements, liens or encumbrances, or claims thereof, not shown by the Public Records.")

add_numbered_item(doc, 5,
    "Any encroachment, encumbrance, violation, variation, or adverse circumstance affecting the Title that "
    "would be disclosed by an accurate and complete land survey of the Land and not shown by the Public Records.")

add_numbered_item(doc, 6,
    "Any mineral or mineral rights leased, granted or retained by current or prior owners.")

p_normal(doc,
    "NOTE: Exceptions Numbered above will be hereby deleted upon issuance of the Loan Policy Only.",
    bold=True, space_before=6, space_after=10)

# Exception 7 — Taxes
add_numbered_item(doc, 7,
    "Taxes and assessments for the year 2026, and subsequent years, not yet due and payable.")

# Exception 8 — MASTER ASSOCIATION (T103/260 lineage)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.first_line_indent = Inches(-0.35)
run_num = p.add_run("8.\t")
run_num.font.size = Pt(10)
run_num.font.bold = True
run_text = p.add_run(
    "Covenants, conditions, restrictions, reservations, easements, liens for assessments, private charges, "
    "option, rights of first refusal, rights of prior approval of future purchaser or occupant, powers of attorney "
    "and limitations on title as to the Declaration of Covenants, Conditions and Restrictions for Rivers Point "
    "Planned-Unit Development, Phase One, recorded in Book T103 at Page 260 on February 11, 1974 in the "
    "official records of the Charleston County Register of Deeds, but omitting any covenant or restriction based "
    "on race, color, religion, sex, handicap, familial status, or national origin unless and only to the extent that "
    "said covenant (a) is exempt under Chapter 42, Section 3607 of the United States Code, or (b) related to "
    "handicap but does not discriminate against handicapped persons; and as amended by that Amendment to "
    "Declaration of CC&Rs and Bylaws \u2014 annexation of condo regime (Reconfiguration Agreement) recorded "
    "in Book J358 at Page 635 on November 9, 2000; and as further amended by that Bylaws of Rivers Point "
    "Homeowners Association, Inc. recorded in Book Z325 at Page 679 on May 7, 1999; and as further amended "
    "by that Amendment to Bylaws \u2014 insurance provisions (Section 10.4) recorded in Book S469 at Page 201 on "
    "October 1, 2003; and as further amended by that Amendment to Bylaws \u2014 insurance provisions (Master Deed "
    "Article IV, Section 3) recorded in Book Z469 at Page 165 on October 1, 2003; and as further amended by that "
    "Third Amendment to Bylaws \u2014 proxy voting provisions (Article 2.4) recorded in Book 0764 at Page 336 on "
    "December 5, 2018; and as further amended by that Certificate of Rules & Regulations \u2014 Rivers Point HOA "
    "(SC Code 27-30-110) recorded in Book 0770 at Page 404 on January 4, 2019; and as further amended by that "
    "Board Resolution \u2014 pool age restriction recorded in Book 0816 at Page 780 on August 15, 2019; and as "
    "further amended by that Fourth Amendment to Bylaws \u2014 special assessment of 0.5% of sales price upon lot "
    "transfer recorded in Book 1227 at Page 021 on February 5, 2024; and as further amended from time to time."
)
run_text.font.size = Pt(10)

# Exception 9 — SUB-ASSOCIATION (A131/001 lineage — lot owners' restrictive covenants)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.first_line_indent = Inches(-0.35)
run_num = p.add_run("9.\t")
run_num.font.size = Pt(10)
run_num.font.bold = True
run_text = p.add_run(
    "Subject to the Restrictive Covenants for Rivers Point Subdivision, executed by all lot owners "
    "(superseding prior covenants in Book T120, Page 135), recorded in Book A131 at Page 001 on March 21, "
    "1983 in the official records of the Charleston County Register of Deeds, but omitting any covenant or "
    "restriction based on race, color, religion, sex, handicap, familial status, or national origin unless and "
    "only to the extent that said covenant (a) is exempt under Chapter 42, Section 3607 of the United States "
    "Code, or (b) related to handicap but does not discriminate against handicapped persons; and as amended "
    "by that Modification of Restrictions \u2014 setback variance for Lot 71, Block B (Architectural Committee) "
    "recorded in Book C302 at Page 247 on May 4, 1998; and as further amended by that Amendment of "
    "Restrictive Covenants \u2014 Architectural Committee election (2005) recorded in Book Y580 at Page 509 on "
    "April 25, 2006; and as further amended by that Amendment of Restrictive Covenants \u2014 extension of "
    "Architectural Committee powers & new representatives recorded in Book O647 at Page 725 on "
    "December 27, 2007; and as further amended by that Amendment of Restrictive Covenants \u2014 Architectural "
    "Committee election (2008) recorded in Book 0045 at Page 349 on April 2, 2009; and as further amended "
    "from time to time."
)
run_text.font.size = Pt(10)

# Exception 10 — SCE&G Easement
add_numbered_item(doc, 10,
    "Subject to the Easements as to South Carolina Electric & Gas Company (SCE&G), recorded in the official "
    "records of the Charleston County Register of Deeds in Book D105 at Page 151 on June 26, 1979. Blanket "
    "utility easement over entire Rivers Point development for electric, gas, pipes, and conduits \u2014 10-foot strip "
    "each side of mains; 15-foot clearance at transformers.")

# Exception 11 — Transfer Fee Covenant
add_numbered_item(doc, 11,
    "Subject to the Notice of Transfer Fee Covenant \u2014 0.5% of sales price general assessment due on each "
    "sale/resale, recorded in Book 1227 at Page 021 on February 5, 2024.")

# Exception 12 — Plat
add_numbered_item(doc, 12,
    "Subject to the Plat recorded in Plat Book AQ at Page 31 (May 1980, C. Roger Jennings, Surveyor) in "
    "the Charleston County Register of Deeds.")

# Exception 13 — Divorce / title defect
add_numbered_item(doc, 13,
    "Title to the subject property is vested in Keith T. Borg and Kathy H. Borg. Per Decree of Divorce entered "
    "August 30, 2017 (Case No. 2017-DR-10-2419, Family Court, 9th Judicial Circuit, Charleston County), "
    "Wife (Kathy Hendrickson Borg) was awarded sole use, possession, and ownership of the marital home. "
    "Husband (Keith Thomas Borg) was ordered to execute a Quit Claim Deed upon presentation by Wife. "
    "No such Quit Claim Deed has been located of record. Both parties must join in any conveyance or, "
    "alternatively, a QCD from Keith Thomas Borg must be recorded prior to closing.")

# Exception 14 — DOR lien
add_numbered_item(doc, 14,
    "SC Department of Revenue State Tax Lien #1479166 against Borg, Keith, filed December 19, 2025, in "
    "the amount of $390.10 (Individual Income Tax for period ending 12/31/2024). Tax $298.00 + Penalty "
    "$39.42 + Interest $34.59 + Costs $18.09. This lien attaches to all real and personal property statewide "
    "and must be satisfied prior to or at closing.")

# Exception 15 — catch-all
add_numbered_item(doc, 15,
    "The Company may make other requirements or exceptions upon its review of the proposed documents "
    "creating the estate or interest to be insured or otherwise ascertaining details of the transaction.")

add_footer_text(doc)

# ============================================================
#  SAVE
# ============================================================
doc.save(OUTPUT)
print(f"Done! Saved to: {OUTPUT}")
