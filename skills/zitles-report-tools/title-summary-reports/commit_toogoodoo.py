#!/usr/bin/env python3
"""Generate a draft ALTA Title Commitment (2021 form) for 7262 Toogoodoo Road."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Draft Commitment - 7262 Toogoodoo Road Hollywood.docx")


def add_page_break(doc):
    doc.add_page_break()


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def add_header_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_right_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(9)


def add_page_header(doc):
    add_right_header(doc, "Commitment for Title Insurance")
    add_right_header(doc, "South Carolina - 2021 v. 01.00 (07-01-2021)")
    add_header_line(doc)


def add_footer_text(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "999999")
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(
        "This page is only a part of a 2021 ALTA Commitment for Title Insurance. "
        "This Commitment is not valid without the Notice; the Commitment to Issue Policy; "
        "the Commitment Conditions; Schedule A; Schedule B, Part I\u2014Requirements; "
        "Schedule B, Part II\u2014Exceptions; and a counter-signature by the Company or "
        "its issuing agent that may be in electronic form."
    )
    run.font.size = Pt(7)
    run.font.italic = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    pPr2 = p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    top2 = OxmlElement("w:top")
    top2.set(qn("w:val"), "single")
    top2.set(qn("w:sz"), "4")
    top2.set(qn("w:space"), "1")
    top2.set(qn("w:color"), "000000")
    pBdr2.append(top2)
    pPr2.append(pBdr2)
    run2 = p2.add_run(
        "Copyright 2021 American Land Title Association. All rights reserved.\n"
        "The use of this Form (or any derivative thereof) is restricted to ALTA licensees and\n"
        "ALTA members in good standing as of the date of use. All other uses are prohibited.\n"
        "Reprinted under license from the American Land Title Association."
    )
    run2.font.size = Pt(7)
    run2.font.bold = True


def p_normal(doc, text, bold=False, size=10, space_after=6, space_before=0, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def p_center_bold(doc, text, size=12, space_before=12, space_after=8):
    return p_normal(doc, text, bold=True, size=size, space_after=space_after,
                    space_before=space_before, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_numbered_item(doc, number, text, bold_number=True, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    run_num = p.add_run(f"{number}.\t")
    run_num.font.size = Pt(10)
    run_num.font.bold = bold_number
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    return p


def add_sub_item(doc, letter, text, indent=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run_letter = p.add_run(f"{letter}.\t")
    run_letter.font.size = Pt(10)
    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    return p


# ============================================================
#  BUILD THE DOCUMENT
# ============================================================
doc = Document()
set_narrow_margins(doc)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)

# ============================================================
#  PAGE 1 — NOTICE / COMMITMENT TO ISSUE POLICY
# ============================================================
add_page_header(doc)

p_center_bold(doc, "ALTA COMMITMENT FOR TITLE INSURANCE\nissued by\n[INSURANCE CARRIER]", size=11, space_before=24)

p_center_bold(doc, "NOTICE", size=11, space_before=18)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run("IMPORTANT\u2014READ CAREFULLY")
run.font.size = Pt(10)
run.font.bold = True
run2 = p.add_run(
    ": THIS COMMITMENT IS AN OFFER TO ISSUE ONE OR MORE TITLE INSURANCE POLICIES. "
    "ALL CLAIMS OR REMEDIES SOUGHT AGAINST THE COMPANY INVOLVING THE CONTENT OF THIS "
    "COMMITMENT OR THE POLICY MUST BE BASED SOLELY IN CONTRACT."
)
run2.font.size = Pt(10)

p_normal(doc,
    "THIS COMMITMENT IS NOT AN ABSTRACT OF TITLE, REPORT OF THE CONDITION OF TITLE, LEGAL "
    "OPINION, OPINION OF TITLE, OR OTHER REPRESENTATION OF THE STATUS OF TITLE. THE "
    "PROCEDURES USED BY THE COMPANY TO DETERMINE INSURABILITY OF THE TITLE, INCLUDING ANY "
    "SEARCH AND EXAMINATION, ARE PROPRIETARY TO THE COMPANY, WERE PERFORMED SOLELY FOR "
    "THE BENEFIT OF THE COMPANY, AND CREATE NO EXTRACONTRACTUAL LIABILITY TO ANY PERSON, "
    "INCLUDING A PROPOSED INSURED.", size=10, space_after=8)

p_normal(doc,
    "THE COMPANY\u2019S OBLIGATION UNDER THIS COMMITMENT IS TO ISSUE A POLICY TO A PROPOSED "
    "INSURED IDENTIFIED IN SCHEDULE A IN ACCORDANCE WITH THE TERMS AND CONDITIONS OF THIS "
    "COMMITMENT. THE COMPANY HAS NO LIABILITY OR OBLIGATION INVOLVING THE CONTENT OF THIS "
    "COMMITMENT TO ANY OTHER PERSON.", size=10, space_after=12)

p_center_bold(doc, "COMMITMENT TO ISSUE POLICY", size=11)

p_normal(doc,
    'Subject to the Notice; Schedule B, Part I\u2014Requirements; Schedule B, Part II\u2014Exceptions; '
    'and the Commitment Conditions, [Insurance Carrier], a [State] Corporation (the \u201cCompany\u201d), '
    'commits to issue the Policy according to the terms and provisions of this Commitment. This '
    'Commitment is effective as of the Commitment Date shown in Schedule A for each Policy described '
    'in Schedule A, only when the Company has entered in Schedule A both the specified dollar amount '
    'as the Proposed Amount of Insurance and the name of the Proposed Insured.', space_after=8)

p_normal(doc,
    "If all of the Schedule B, Part I\u2014Requirements have not been met within 180 days after the "
    "Commitment Date, this Commitment terminates and the Company\u2019s liability and obligation end.",
    space_after=12)

p_normal(doc, "[INSURANCE CARRIER]", bold=True, space_after=18)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("By: ____________________________\t\t\tBy: ____________________________")
run.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("      [Name], President\t\t\t\t\t      [Name], Secretary")
run2.font.size = Pt(9)
run2.font.bold = True

add_footer_text(doc)

# ============================================================
#  PAGE 2-4 — COMMITMENT CONDITIONS
# ============================================================
add_page_break(doc)
add_page_header(doc)
p_center_bold(doc, "COMMITMENT CONDITIONS", size=11, space_before=6)

add_numbered_item(doc, 1, "DEFINITIONS")
definitions = [
    ("a", '"Discriminatory Covenant": Any covenant, condition, restriction, or limitation that is unenforceable under applicable law because it illegally discriminates against a class of individuals based on personal characteristics such as race, color, religion, sex, sexual orientation, gender identity, familial status, disability, national origin, or other legally protected class.'),
    ("b", '"Knowledge" or "Known": Actual knowledge or actual notice, but not constructive notice imparted by the Public Records.'),
    ("c", '"Land": The land described in Item 5 of Schedule A and improvements located on that land that by State law constitute real property. The term "Land" does not include any property beyond that described in Schedule A, nor any right, title, interest, estate, or easement in any abutting street, road, avenue, alley, lane, right-of-way, body of water, or waterway, but does not modify or limit a right of access to and from the Land is to be insured by the Policy.'),
    ("d", '"Mortgage": A mortgage, deed of trust, trust deed, security deed, or other real property security instrument, including one evidenced by electronic means authorized by law.'),
    ("e", '"Policy": Each contract of title insurance, in a form adopted by the American Land Title Association, issued or to be issued by the Company pursuant to this Commitment.'),
    ("f", '"Proposed Amount of Insurance": Each dollar amount specified in Schedule A as the Proposed Amount of Insurance of each Policy to be issued pursuant to this Commitment.'),
    ("g", '"Proposed Insured": Each person identified in Schedule A as the Proposed Insured of each Policy to be issued pursuant to this Commitment.'),
    ("h", '"Public Records": The recording or filing system established under State statutes in effect at the Commitment Date under which a document must be recorded or filed to impart constructive notice of matters relating to the Title to a purchaser for value without Knowledge. The term "Public Records" does not include any other recording or filing system, including any pertaining to environmental remediation or protection, planning, permitting, zoning, licensing, building, health, public safety, or national security matters.'),
    ("i", '"State": The state or commonwealth of the United States within whose exterior boundaries the Land is located. The term "State" also includes the District of Columbia, the Commonwealth of Puerto Rico, the U.S. Virgin Islands, and Guam.'),
    ("j", '"Title": The estate or interest in the Land identified in Item 3 of Schedule A.'),
]
for letter, text in definitions:
    add_sub_item(doc, letter, text)

add_numbered_item(doc, 2,
    "If all of the Schedule B, Part I\u2014Requirements have not been met within the time period specified "
    "in the Commitment to Issue Policy, this Commitment terminates and the Company\u2019s liability and obligation end.")

add_numbered_item(doc, 3, "The Company\u2019s liability and obligation is limited by and this Commitment is not valid without:")
for letter, text in [("a", "the Notice;"), ("b", "the Commitment to Issue Policy;"),
                      ("c", "the Commitment Conditions;"), ("d", "Schedule A;"),
                      ("e", "Schedule B, Part I\u2014Requirements;"),
                      ("f", "Schedule B, Part II\u2014Exceptions; and"),
                      ("g", "a counter-signature by the Company or its issuing agent that may be in electronic form.")]:
    add_sub_item(doc, letter, text)

add_page_break(doc)
add_page_header(doc)

add_numbered_item(doc, 4, "COMPANY\u2019S RIGHT TO AMEND")
p_normal(doc,
    "The Company may amend this Commitment at any time. If the Company amends this Commitment to add "
    "a defect, lien, encumbrance, adverse claim, or other matter recorded in the Public Records prior to the "
    "Commitment Date, any liability of the Company is limited by Commitment Condition 5. The Company is not "
    "liable for any other amendment to this Commitment.", space_after=8)

add_numbered_item(doc, 5, "LIMITATIONS OF LIABILITY")
add_sub_item(doc, "a",
    "The Company\u2019s liability under Commitment Condition 4 is limited to the Proposed Insured\u2019s actual "
    "expense incurred in the interval between the Company\u2019s delivery to the Proposed Insured of the "
    "Commitment and the delivery of the amended Commitment, resulting from the Proposed Insured\u2019s good "
    "faith reliance to:")
for sub in ["comply with the Schedule B, Part I\u2014Requirements;",
            "eliminate, with the Company\u2019s written consent, any Schedule B, Part II\u2014Exceptions; or",
            "acquire the Title or create the Mortgage covered by this Commitment."]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(1.5)
    run = p.add_run(sub)
    run.font.size = Pt(10)

add_sub_item(doc, "b", "The Company is not liable under Commitment Condition 5.a. if the Proposed Insured requested the amendment or had Knowledge of the matter and did not notify the Company about it in writing.")
add_sub_item(doc, "c", "The Company is only liable under Commitment Condition 4 if the Proposed Insured would not have incurred the expense had the Commitment included the added matter when the Commitment was first delivered to the Proposed Insured.")
add_sub_item(doc, "d", "The Company\u2019s liability does not exceed the lesser of the Proposed Insured\u2019s actual expense incurred in good faith and described in Commitment Condition 5.a. or the Proposed Amount of Insurance.")
add_sub_item(doc, "e", "The Company is not liable for the content of the Transaction Identification Data, if any.")
add_sub_item(doc, "f", "The Company is not obligated to issue the Policy referred to in this Commitment unless all of the Schedule B, Part I\u2014Requirements have been met to the satisfaction of the Company.")
add_sub_item(doc, "g", "The Company\u2019s liability is further limited by the terms and provisions of the Policy to be issued to the Proposed Insured.")

add_numbered_item(doc, 6, "LIABILITY OF THE COMPANY MUST BE BASED ON THIS COMMITMENT; CHOICE OF LAW AND CHOICE OF FORUM")
add_sub_item(doc, "a", "Only a Proposed Insured identified in Schedule A, and no other person, may make a claim under this Commitment.")
add_sub_item(doc, "b", "Any claim must be based in contract under the State law of the State where the Land is located and is restricted to the terms and provisions of this Commitment. Any litigation or other proceeding brought by the Proposed Insured against the Company must be filed only in a State or federal court having jurisdiction.")
add_sub_item(doc, "c", "This Commitment, as last revised, is the exclusive and entire agreement between the parties with respect to the subject matter of this Commitment and supersedes all prior commitment negotiations, representations, and proposals of any kind, whether written or oral, express or implied, relating to the subject matter of this Commitment.")
add_sub_item(doc, "d", "The deletion or modification of any Schedule B, Part II\u2014Exception does not constitute an agreement or obligation to provide coverage beyond the terms and provisions of this Commitment or the Policy.")
add_sub_item(doc, "e", "Any amendment or endorsement to this Commitment must be in writing and authenticated by a person authorized by the Company.")
add_sub_item(doc, "f", "When the Policy is issued, all liability and obligation under this Commitment will end and the Company\u2019s only liability will be under the Policy.")

add_page_break(doc)
add_page_header(doc)

add_numbered_item(doc, 7, "IF THIS COMMITMENT IS ISSUED BY AN ISSUING AGENT")
p_normal(doc, "The issuing agent is the Company\u2019s agent only for the limited purpose of issuing title insurance commitments and policies. The issuing agent is not the Company\u2019s agent for closing, settlement, escrow, or any other purpose.", space_after=8)

add_numbered_item(doc, 8, "PRO-FORMA POLICY")
p_normal(doc, "The Company may provide, at the request of a Proposed Insured, a pro-forma policy illustrating the coverage that the Company may provide. A pro-forma policy neither reflects the status of Title at the time that the pro-forma policy is delivered to a Proposed Insured, nor is it a commitment to insure.", space_after=8)

add_numbered_item(doc, 9, "CLAIMS PROCEDURES")
p_normal(doc, "This Commitment incorporates by reference all Conditions for making a claim in the Policy to be issued to the Proposed Insured. Commitment Condition 9 does not modify the limitations of liability in Commitment Conditions 5 and 6.", space_after=8)

add_numbered_item(doc, 10, "CLASS ACTION")
p_normal(doc, "ALL CLAIMS AND DISPUTES ARISING OUT OF OR RELATING TO THIS COMMITMENT, INCLUDING ANY SERVICE OR OTHER MATTER IN CONNECTION WITH ISSUING THIS COMMITMENT, ANY BREACH OF A COMMITMENT PROVISION, OR ANY OTHER CLAIM OR DISPUTE ARISING OUT OF OR RELATING TO THE TRANSACTION GIVING RISE TO THIS COMMITMENT, MUST BE BROUGHT IN AN INDIVIDUAL CAPACITY. NO PARTY MAY SERVE AS PLAINTIFF, CLASS MEMBER, OR PARTICIPANT IN ANY CLASS OR REPRESENTATIVE PROCEEDING. ANY POLICY ISSUED PURSUANT TO THIS COMMITMENT WILL CONTAIN A CLASS ACTION CONDITION.", size=10, space_after=8)

add_numbered_item(doc, 11, "ARBITRATION")
p_normal(doc, "The Policy contains an arbitration clause. All arbitrable matters when the Proposed Amount of Insurance is $2,000,000 or less may be arbitrated at the election of either the Company or the Proposed Insured as the exclusive remedy of the parties. A Proposed Insured may review a copy of the arbitration rules at http://www.alta.org/arbitration.", space_after=8)

add_footer_text(doc)

# ============================================================
#  TRANSACTION ID + SCHEDULE A
# ============================================================
add_page_break(doc)
add_page_header(doc)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Transaction Identification Data, for which the Company assumes no liability as set forth in Commitment Condition 5.e.:")
run.font.size = Pt(10)
run.font.bold = True

tid_lines = [
    ("Issuing Agent:", "[Title Agency Name]"),
    ("Issuing Office:", "[Address]"),
    ("", "[City, State ZIP]"),
    ("Issuing Office\u2019s ALTA\u00ae Registry ID:", ""),
    ("Loan ID Number:", ""),
    ("Commitment Number:", "[File Number]"),
    ("Issuing Office File Number:", "[File Number]"),
    ("Property Address:", "7262 Toogoodoo Road, Hollywood, SC 29449"),
    ("Revision Number:", ""),
]
for label, value in tid_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if label:
        run_l = p.add_run(f"{label} ")
        run_l.font.size = Pt(10)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)

p_center_bold(doc, "SCHEDULE A", size=12, space_before=18)

add_numbered_item(doc, 1, "Commitment Date:  at 8:00 AM")

add_numbered_item(doc, 2, "Policy to be issued:")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.left_indent = Inches(1.0)
lines = [
    ("a.\t2021 ALTA Homeowner\u2019s Policy", ""),
    ("\tProposed Insured:", "\t\t[Proposed Insured]"),
    ("\tProposed Amount of Insurance:", "\t$___________"),
    ("\tThe estate or interest to be insured:", "\tfee simple"),
]
for left, right in lines:
    run = p.add_run(f"{left}{right}\n")
    run.font.size = Pt(10)

add_numbered_item(doc, 3, "The estate or interest in the Land at the Commitment Date is:\n\tfee simple")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.first_line_indent = Inches(-0.35)
run_num = p.add_run("4.\t")
run_num.font.size = Pt(10)
run_num.font.bold = True
run_text = p.add_run(
    "The Title is, at the Commitment Date, vested in:\n"
    "\tTracy E. Miller and Roanna B. Miller, as Joint Tenants with Rights of Survivorship "
    "and Not as Tenants in Common, by deed from M. Scott Villas, Trustee of the M. Scott "
    "Villas Revocable Living Trust, UTD 07/23/2021, and Heather K. Villas, Trustee of the "
    "Heather K. Villas Revocable Living Trust, UTD 07/23/2021, dated 06/11/2024 and recorded "
    "with the Charleston County Register of Deeds on 07/26/2024 in Book 1258, Page 826."
)
run_text.font.size = Pt(10)

add_numbered_item(doc, 5, "The Land is described as follows:")
p_normal(doc, "SEE EXHIBIT A ATTACHED HERETO AND MADE A PART HEREOF",
         bold=True, space_before=6, space_after=18)

# Title Agency block
p_normal(doc, "[TITLE AGENCY NAME]", bold=True, space_after=2)
p_normal(doc, "[Address]\n[City, State ZIP]\nTelephone:", space_after=12)
p_normal(doc, "Countersigned:", space_after=12)

p = doc.add_paragraph()
run = p.add_run("By:____________________________")
run.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run("      Authorized Signatory")
run2.font.size = Pt(10)
run2.font.bold = True

p3 = doc.add_paragraph()
run3 = p3.add_run("[Name], License #[______]\n[Title Agency], License #")
run3.font.size = Pt(10)

add_footer_text(doc)

# ============================================================
#  EXHIBIT A
# ============================================================
add_page_break(doc)
add_page_header(doc)

p_center_bold(doc, "EXHIBIT A", size=12, space_before=6)

p_normal(doc, "Commitment No.: [File Number]", bold=True, space_after=12)

p_normal(doc,
    "The land referred to herein below is situated in the County of Charleston, State of South Carolina, "
    "and is described as follows:", space_after=12)

p_normal(doc,
    "ALL that certain piece, parcel and lot of land, together with the buildings and improvements thereon, "
    "situate, lying and being in St. Paul\u2019s Parish, Charleston County, South Carolina. Measuring and "
    "containing five (5) acres of high land and approximately eleven (11) acres of marsh land.",
    space_after=8)

p_normal(doc,
    "BUTTING AND BOUNDING to the North and Northeast on a spur of the Atlantic Coast Line Railroad, "
    "to the East and Southeast on State Highway No. 6, on the West Southwest, West and Northwest on "
    "Swinton Creek.",
    space_after=8)

p_normal(doc,
    "The property herein conveyed being more particularly described in a plat prepared by John McCrady "
    "Company, engineers dated March 1946, and attached to a deed dated January 16, 1947 and recorded "
    "January 23, 1947 in Book C-45, at Page 522 in the ROD Office for Charleston County.",
    space_after=8)

p_normal(doc, "TMS: 164-00-00-234", space_after=8)

p_normal(doc,
    "This being the same property conveyed to Tracy E. Miller and Roanna B. Miller, as Joint Tenants "
    "with Rights of Survivorship and Not as Tenants in Common, by deed of M. Scott Villas, Trustee of "
    "the M. Scott Villas Revocable Living Trust, UTD 07/23/2021, and Heather K. Villas, Trustee of the "
    "Heather K. Villas Revocable Living Trust, UTD 07/23/2021, dated June 11, 2024, and recorded at the "
    "Charleston County ROD Office on July 26, 2024, in Book 1258, at Page 826.",
    space_after=8)

add_footer_text(doc)

# ============================================================
#  SCHEDULE B, PART I — REQUIREMENTS
# ============================================================
add_page_break(doc)
add_page_header(doc)

p_center_bold(doc, "SCHEDULE B, PART I\u2014Requirements", size=12, space_before=6)

p_normal(doc, "All of the following Requirements must be met:", space_after=8)

add_numbered_item(doc, 1,
    "The Proposed Insured must notify the Company in writing of the name of any party not referred to in this "
    "Commitment who will obtain an interest in the Land or who will make a loan on the Land. The Company "
    "may then make additional Requirements or Exceptions.")

add_numbered_item(doc, 2, "Pay the agreed amount for the estate or interest to be insured.")

add_numbered_item(doc, 3, "Pay the premiums, fees, and charges for the Policy to the Company.")

add_numbered_item(doc, 4,
    "We must be furnished with a copy of SCID 3601 executed pursuant to Section 38-75-960 S.C. Code of "
    "Laws 1976, as amended, and an executed Notice of Availability of Title Insurance pursuant to S.C. "
    "Insurance Department Regulation R-69-18, Vol. 25A of S.C. Code of Laws 1976, as amended.")

add_numbered_item(doc, 5,
    "Seller\u2019s/Owner\u2019s Affidavit Indemnity executed by current owner(s) of the land on a form to be supplied by "
    "the Company stating that there have been no improvements to the land within the past 90 days which "
    "could give rise to a construction lien and that there are no accounts or claims pending and unpaid which "
    "could constitute a lien against the land.")

add_numbered_item(doc, 6, "Receipt of the acknowledged [Insurance Carrier] Privacy Policy.")

add_numbered_item(doc, 7,
    "Documents satisfactory to the Company that convey the Title or create the Mortgage to be insured, or "
    "both, must be properly authorized, executed, delivered, and recorded in the Public Records.")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run(
    "Deed from Tracy E. Miller and Roanna B. Miller to [Purchaser] conveying the land described "
    "in Schedule A herein.")
run.font.size = Pt(10)

add_numbered_item(doc, 8,
    "Payoff, satisfaction and release of the mortgage recorded in Book 1258, Page 830 from Tracy E. Miller "
    "and Roanna B. Miller to Intercoastal Mortgage, LLC (MERS as nominee), in the amount of $567,920.00, "
    "recorded on July 26, 2024 in the Office of the Charleston County Register of Deeds.")

add_numbered_item(doc, 9,
    "Satisfaction or release of the seller purchase money mortgage from Pauline Holden to Bruce Weitz and "
    "Marlene J. Weitz, in the amount of $33,500.00, recorded on March 21, 1997 in the Office of the "
    "Charleston County Register of Deeds in Book O-281, Page 410. NOTE: This mortgage matured on "
    "March 20, 2002 and no satisfaction has been recorded. The property has conveyed five (5) times since "
    "this mortgage was recorded without a recorded release (Holden \u2192 Jordan \u2192 Curry \u2192 Villas/Wilkinson "
    "\u2192 Villas Trust \u2192 Miller).")

add_numbered_item(doc, 10,
    "Resolution of the pending eminent domain condemnation action filed by the South Carolina Department "
    "of Transportation, Case No. 2025-CP-10-02869 (Court of Common Pleas, Charleston County), filed "
    "May 16, 2025, seeking to acquire 11,723 square feet (0.27 acre) in fee simple from the subject property "
    "for the S-390 (Toogoodoo Road) Bridge Replacement over Swinton Creek, Project P030449, Tract 11. "
    "Condemnation tender: $3,838.00. The landowners have demanded a jury trial (filed July 10, 2025). "
    "Intercoastal Mortgage, LLC is named as Other Condemnee.")

add_numbered_item(doc, 11,
    "The Company may make other requirements or exceptions upon its review of the proposed documents "
    "creating the estate or interest to be insured or otherwise ascertaining details of the transaction.")

add_footer_text(doc)

# ============================================================
#  SCHEDULE B, PART II — EXCEPTIONS
# ============================================================
add_page_break(doc)
add_page_header(doc)

p_center_bold(doc, "SCHEDULE B, PART II\u2014Exceptions", size=12, space_before=6)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run(
    "Some historical land records contain Discriminatory Covenants that are illegal and unenforceable by law. "
    "This Commitment and the Policy treat any Discriminatory Covenant in a document referenced in Schedule "
    "B as if each Discriminatory Covenant is redacted, repudiated, removed, and not republished or recirculated. "
    "Only the remaining provisions of the document will be excepted from coverage."
)
run.font.size = Pt(10)
run.font.bold = True

p_normal(doc,
    "The Policy will not insure against loss or damage resulting from the terms and conditions of any lease or easement "
    "identified in Schedule A, and will include the following Exceptions unless cleared to the satisfaction of the Company:",
    space_after=8)

# Standard exceptions 1-6
add_numbered_item(doc, 1,
    "Any defect, lien, encumbrance, adverse claim, or other matter that appears for the first time in the Public "
    "Records or is created, attaches, or is disclosed between the Commitment Date and the date on which all of "
    "the Schedule B, Part I\u2014Requirements are met.")

add_numbered_item(doc, 2,
    "(a) Taxes or assessments that are not shown as existing liens by the records of any taxing authority that "
    "levies taxes or assessments on real property or by the Public Records; (b) proceedings by a public agency "
    "that may result in taxes or assessments, or notices of such proceedings, whether or not shown by the "
    "records of such agency or by the Public Records.")

add_numbered_item(doc, 3,
    "Any facts, rights, interests, or claims that are not shown by the Public Records but that could be "
    "ascertained by an inspection of the Land or that may be asserted by persons in possession of the Land.")

add_numbered_item(doc, 4, "Easements, liens or encumbrances, or claims thereof, not shown by the Public Records.")

add_numbered_item(doc, 5,
    "Any encroachment, encumbrance, violation, variation, or adverse circumstance affecting the Title that "
    "would be disclosed by an accurate and complete land survey of the Land and not shown by the Public Records.")

add_numbered_item(doc, 6, "Any mineral or mineral rights leased, granted or retained by current or prior owners.")

p_normal(doc,
    "NOTE: Exceptions Numbered above will be hereby deleted upon issuance of the Loan Policy Only.",
    bold=True, space_before=6, space_after=10)

# Exception 7 — Taxes
add_numbered_item(doc, 7, "Taxes and assessments for the year 2026, and subsequent years, not yet due and payable.")

# Exception 8 — SCE&G Easement
add_numbered_item(doc, 8,
    "Subject to the Right of Way Grant from Pauline Holden to South Carolina Electric & Gas Company "
    "(SCE&G), dated November 20, 2003, recorded December 30, 2003 in Book D480 at Page 167 in the "
    "official records of the Charleston County Register of Deeds, granting permission to install, maintain, "
    "and extend a guy anchor and wire at a point near the property\u2019s southern boundary on the north side "
    "of Toogoodoo Road.")

# Exception 9 — Plat
add_numbered_item(doc, 9,
    "Subject to the Plat of property containing 5 acres of high land and approximately 11 acres of marsh, "
    "situated in St. Paul\u2019s Parish, Charleston County, S.C., surveyed for the Charleston Council of Girl "
    "Scouts by The John McCrady Company, Engineers, dated March 1946, recorded in Book C-45, at Page "
    "522 in the Charleston County ROD Office.")

# Exception 10 — Pending Condemnation
add_numbered_item(doc, 10,
    "Pending eminent domain condemnation action: South Carolina Department of Transportation v. Tracy E. "
    "Miller and Roanna B. Miller, as Joint Tenants with Rights of Survivorship, and Intercoastal Mortgage, "
    "LLC, Mortgagee, Case No. 2025-CP-10-02869, Court of Common Pleas, Charleston County, filed "
    "May 16, 2025. SCDOT seeks to acquire 11,723 square feet (0.27 acre) in fee simple for the construction "
    "of RTE./RD.S-390 (Toogoodoo Road) Bridge Replacement over Swinton Creek, Project ID P030449, "
    "Tract 11. Lis Pendens, Summons, and Condemnation Notice and Tender of Payment ($3,838.00) filed. "
    "Landowners have demanded a jury trial (filed July 10, 2025, attorney John Edward Robinson, "
    "SC Bar No. 75919, 36 Broad Street, Charleston, SC 29401).")

# Exception 11 — Open Mortgage
add_numbered_item(doc, 11,
    "Subject to the Mortgage to Real Estate from Pauline Holden (mortgagor) to Bruce Weitz and Marlene J. "
    "Weitz (mortgagee), in the penal sum of $33,500.00 at 8.0% interest, dated March 20, 1997, recorded "
    "March 21, 1997 in Book O-281 at Page 410 in the official records of the Charleston County Register of "
    "Deeds. This is a seller purchase money mortgage, second and subordinate to Pauline Holden\u2019s mortgage "
    "to NationsBank, N.A. (Carolinas) recorded simultaneously. This mortgage matured on March 20, 2002. "
    "No satisfaction or release has been located of record. The property has conveyed five (5) times since "
    "without a recorded satisfaction.")

# Exception 12 — catch-all
add_numbered_item(doc, 12,
    "The Company may make other requirements or exceptions upon its review of the proposed documents "
    "creating the estate or interest to be insured or otherwise ascertaining details of the transaction.")

add_footer_text(doc)

# ============================================================
#  SAVE
# ============================================================
doc.save(OUTPUT)
print(f"Done! Saved to: {OUTPUT}")
