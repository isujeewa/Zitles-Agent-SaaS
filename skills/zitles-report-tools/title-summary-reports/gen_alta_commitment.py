#!/usr/bin/env python3
"""Generate an ALTA-format Draft Title Commitment as a DOCX."""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_para(doc, text, bold=False, size=10, align=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p


def build(data, output_path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ---- COVER ----
    add_para(doc, "DRAFT — ALTA COMMITMENT FOR TITLE INSURANCE", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "issued by\n[TITLE INSURANCE COMPANY]", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_para(doc, "NOTICE", bold=True, size=10)
    add_para(doc, 'IMPORTANT\u2014READ CAREFULLY: THIS COMMITMENT IS AN OFFER TO ISSUE ONE OR MORE TITLE INSURANCE POLICIES. ALL CLAIMS OR REMEDIES SOUGHT AGAINST THE COMPANY INVOLVING THE CONTENT OF THIS COMMITMENT OR THE POLICY MUST BE BASED SOLELY IN CONTRACT.', size=8)
    add_para(doc, 'THIS COMMITMENT IS NOT AN ABSTRACT OF TITLE, REPORT OF THE CONDITION OF TITLE, LEGAL OPINION, OPINION OF TITLE, OR OTHER REPRESENTATION OF THE STATUS OF TITLE. THE PROCEDURES USED BY THE COMPANY TO DETERMINE INSURABILITY OF THE TITLE, INCLUDING ANY SEARCH AND EXAMINATION, ARE PROPRIETARY TO THE COMPANY, WERE PERFORMED SOLELY FOR THE BENEFIT OF THE COMPANY, AND CREATE NO EXTRACONTRACTUAL LIABILITY TO ANY PERSON, INCLUDING A PROPOSED INSURED.', size=8)
    add_para(doc, "THE COMPANY'S OBLIGATION UNDER THIS COMMITMENT IS TO ISSUE A POLICY TO A PROPOSED INSURED IDENTIFIED IN SCHEDULE A IN ACCORDANCE WITH THE TERMS AND PROVISIONS OF THIS COMMITMENT. THE COMPANY HAS NO LIABILITY OR OBLIGATION INVOLVING THE CONTENT OF THIS COMMITMENT TO ANY OTHER PERSON.", size=8, space_after=Pt(12))

    add_para(doc, "COMMITMENT TO ISSUE POLICY", bold=True, size=10)
    add_para(doc, 'Subject to the Notice; Schedule B, Part I\u2014Requirements; Schedule B, Part II\u2014Exceptions; and the Commitment Conditions, [Title Insurance Company] (the "Company"), commits to issue the Policy according to the terms and provisions of this Commitment. This Commitment is effective as of the Commitment Date shown in Schedule A for each Policy described in Schedule A, only when the Company has entered in Schedule A the required items. The Company may amend or supplement this Commitment at any time prior to the Policy effective date.', size=9)
    add_para(doc, "If all of the Schedule B, Part I\u2014Requirements have not been met within 180 days after the Commitment Date, this Commitment terminates and the Company's liability and obligation end.", size=9, space_after=Pt(8))

    add_para(doc, "[TITLE INSURANCE COMPANY]", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_para(doc, "By: ____________________________          By: ____________________________", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "      President                                                    Secretary", size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ---- Transaction ID ----
    add_para(doc, "Transaction Identification Data, for which the Company assumes no liability:", bold=True, size=9, space_after=Pt(2))
    for label, val in [
        ("Issuing Agent:", ""),
        ("Issuing Office:", ""),
        ("Commitment Number:", "_______________"),
        ("Issuing Office File Number:", "_______________"),
        ("Property Address:", data.get("property_address_full", "")),
        ("Revision Number:", ""),
    ]:
        add_para(doc, f"{label} {val}", size=9, space_after=Pt(1))

    doc.add_page_break()

    # ---- SCHEDULE A ----
    add_para(doc, "SCHEDULE A", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_para(doc, f'1.\tCommitment Date: {data.get("commitment_date", "")} at 8:00 AM', size=10)
    add_para(doc, "2.\tPolicy to be issued:", size=10)
    add_para(doc, f'\ta.\t{data.get("policy_type", "2021 ALTA Owner\'s Policy")}', size=10)
    add_para(doc, f'\tProposed Insured:\t\t_______________', size=10)
    add_para(doc, f'\tProposed Amount of Insurance:\t$_______________', size=10)
    add_para(doc, f'\tThe estate or interest to be insured:\tfee simple', size=10, space_after=Pt(8))

    add_para(doc, "3.\tThe estate or interest in the Land at the Commitment Date is:\n\tfee simple", size=10, space_after=Pt(8))

    add_para(doc, "4.\tThe Title is, at the Commitment Date, vested in:", size=10)
    add_para(doc, data.get("vesting_paragraph", ""), size=10, space_after=Pt(8))

    if data.get("prior_chain_paragraph"):
        add_para(doc, data["prior_chain_paragraph"], size=10, space_after=Pt(8))

    add_para(doc, "5.\tThe Land is described as follows:", size=10, space_after=Pt(4))
    add_para(doc, "SEE EXHIBIT A ATTACHED HERETO AND MADE A PART HEREOF", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

    add_para(doc, "[ISSUING AGENT]", bold=True, size=10)
    add_para(doc, "Telephone:", size=9)
    add_para(doc, "Countersigned:", size=9, space_after=Pt(8))
    add_para(doc, "By: ____________________________", size=10)
    add_para(doc, "      Authorized Signatory", size=9)

    doc.add_page_break()

    # ---- EXHIBIT A ----
    add_para(doc, "EXHIBIT A", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_para(doc, "Commitment No.: _______________", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_para(doc, f'The land referred to herein below is situated in the County of {data.get("county", "")}, State of South Carolina, and is described as follows:', size=10, space_after=Pt(8))

    for para_text in data.get("legal_description_paragraphs", []):
        add_para(doc, para_text, size=10, space_after=Pt(6))

    add_para(doc, f'TMS: {data.get("parcel_id", "")}', size=10, space_after=Pt(8))

    doc.add_page_break()

    # ---- SCHEDULE B, PART I ----
    add_para(doc, "SCHEDULE B, PART I\u2014Requirements", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_para(doc, "All of the following Requirements must be met:", size=10, space_after=Pt(6))

    for i, req in enumerate(data.get("requirements", []), 1):
        add_para(doc, f"{i}.\t{req}", size=10, space_after=Pt(4))

    doc.add_page_break()

    # ---- SCHEDULE B, PART II ----
    add_para(doc, "SCHEDULE B, PART II\u2014Exceptions", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    add_para(doc, "Some historical land records contain Discriminatory Covenants that are illegal and unenforceable by law. This Commitment and the Policy treat any Discriminatory Covenant in a document referenced in Schedule B as if each Discriminatory Covenant is redacted, repudiated, removed, and not republished or recirculated.", size=8, space_after=Pt(8))
    add_para(doc, "The Policy will not insure against loss or damage resulting from the terms and conditions of any lease or easement identified in Schedule A, and will include the following Exceptions unless cleared to the satisfaction of the Company:", size=10, space_after=Pt(6))

    for i, exc in enumerate(data.get("exceptions", []), 1):
        add_para(doc, f"{i}.\t{exc}", size=10, space_after=Pt(4))

    doc.save(output_path)
    print(f"Done! DOCX: {output_path}")


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("data_file")
    parser.add_argument("-o", "--output", required=True)
    args = parser.parse_args()
    with open(args.data_file) as f:
        data = json.load(f)
    build(data, args.output)
