#!/usr/bin/env python3
"""Generate a draft ALTA Title Commitment for 236 Senior Road, Kingstree, SC (Williamsburg County)."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Draft Commitment - 236 Senior Road Kingstree.docx")

HIGHLIGHT_YELLOW = 7


def add_header_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_right_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Commitment for Title Insurance\nSouth Carolina - 2021 v. 01.00 (07-01-2021)")
    run.font.size = Pt(9)
    add_header_line(doc)


def add_footer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(
        "This page is only a part of a 2021 ALTA Commitment for Title Insurance. "
        "This Commitment is not valid without the Notice; the Commitment to Issue Policy; "
        "the Commitment Conditions; Schedule A; Schedule B, Part I\u2014Requirements; "
        "Schedule B, Part II\u2014Exceptions; and a counter-signature by the Company or "
        "its issuing agent that may be in electronic form."
    )
    run.font.size = Pt(7)
    run.italic = True

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Copyright 2021 American Land Title Association. All rights reserved.\n"
        "The use of this Form (or any derivative thereof) is restricted to ALTA licensees and\n"
        "ALTA members in good standing as of the date of use. All other uses are prohibited.\n"
        "Reprinted under license from the American Land Title Association.\n"
        "Form 50133045 (7-5-22)"
    )
    run2.font.size = Pt(7)


def hl(p, text):
    """Add highlighted run to paragraph."""
    r = p.add_run(text)
    r.font.highlight_color = HIGHLIGHT_YELLOW
    return r


def centered_bold(doc, text, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def numbered(doc, num, text, bold_num=True):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}.")
    r.bold = bold_num
    p.add_run(f"\t{text}")
    return p


def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===================== PAGE 1 - COVER =====================
    add_right_header(doc)

    centered_bold(doc, "ALTA COMMITMENT FOR TITLE INSURANCE", 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("issued by")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hl(p, "[UNDERWRITER NAME]").bold = True

    doc.add_paragraph()
    centered_bold(doc, "NOTICE", 11)

    p = doc.add_paragraph()
    r = p.add_run("IMPORTANT\u2014READ CAREFULLY")
    r.bold = True
    p.add_run(
        ": THIS COMMITMENT IS AN OFFER TO ISSUE ONE OR MORE TITLE INSURANCE POLICIES. "
        "ALL CLAIMS OR REMEDIES SOUGHT AGAINST THE COMPANY INVOLVING THE CONTENT OF THIS "
        "COMMITMENT OR THE POLICY MUST BE BASED SOLELY IN CONTRACT."
    )

    doc.add_paragraph(
        "THIS COMMITMENT IS NOT AN ABSTRACT OF TITLE, REPORT OF THE CONDITION OF TITLE, "
        "LEGAL OPINION, OPINION OF TITLE, OR OTHER REPRESENTATION OF THE STATUS OF TITLE. "
        "THE PROCEDURES USED BY THE COMPANY TO DETERMINE INSURABILITY OF THE TITLE, INCLUDING "
        "ANY SEARCH AND EXAMINATION, ARE PROPRIETARY TO THE COMPANY, WERE PERFORMED SOLELY FOR "
        "THE BENEFIT OF THE COMPANY, AND CREATE NO EXTRACONTRACTUAL LIABILITY TO ANY PERSON, "
        "INCLUDING A PROPOSED INSURED."
    )

    doc.add_paragraph(
        "THE COMPANY\u2019S OBLIGATION UNDER THIS COMMITMENT IS TO ISSUE A POLICY TO A PROPOSED "
        "INSURED IDENTIFIED IN SCHEDULE A IN ACCORDANCE WITH THE TERMS AND PROVISIONS OF THIS "
        "COMMITMENT. THE COMPANY HAS NO LIABILITY OR OBLIGATION INVOLVING THE CONTENT OF THIS "
        "COMMITMENT TO ANY OTHER PERSON."
    )

    doc.add_paragraph()
    centered_bold(doc, "COMMITMENT TO ISSUE POLICY", 11)

    p = doc.add_paragraph()
    p.add_run(
        "Subject to the Notice; Schedule B, Part I\u2014Requirements; Schedule B, Part II\u2014Exceptions; "
        "and the Commitment Conditions, "
    )
    hl(p, "[UNDERWRITER NAME]")
    p.add_run(
        ", a "
    )
    hl(p, "[State]")
    p.add_run(
        " Corporation (the \u201cCompany\u201d), commits to issue "
        "the Policy according to the terms and provisions of this Commitment. This Commitment is effective "
        "as of the Commitment Date shown in Schedule A for each Policy described in Schedule A, only when "
        "the Company has entered in Schedule A both the specified dollar amount as the Proposed Amount of "
        "Insurance and the name of the Proposed Insured."
    )

    doc.add_paragraph(
        "If all of the Schedule B, Part I\u2014Requirements have not been met within 180 days after "
        "the Commitment Date, this Commitment terminates and the Company\u2019s liability and obligation end."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    hl(p, "[UNDERWRITER NAME]").bold = True

    doc.add_paragraph()
    doc.add_paragraph("By: ____________________________          By: ____________________________")
    p = doc.add_paragraph("       ")
    hl(p, "[Name, Title]")
    p.add_run("                                        ")
    hl(p, "[Name, Title]")

    add_footer(doc)
    doc.add_page_break()

    # ===================== PAGES 2-4 - CONDITIONS =====================
    add_right_header(doc)
    centered_bold(doc, "COMMITMENT CONDITIONS", 11)

    p = doc.add_paragraph()
    r = p.add_run("1.\tDEFINITIONS")
    r.bold = True

    defs = [
        ("a.", "\u201cDiscriminatory Covenant\u201d: Any covenant, condition, restriction, or limitation that is unenforceable under applicable law because it illegally discriminates against a class of individuals based on personal characteristics such as race, color, religion, sex, sexual orientation, gender identity, familial status, disability, national origin, or other legally protected class."),
        ("b.", "\u201cKnowledge\u201d or \u201cKnown\u201d: Actual knowledge or actual notice, but not constructive notice imparted by the Public Records."),
        ("c.", "\u201cLand\u201d: The land described in Item 5 of Schedule A and improvements located on that land that by State law constitute real property. The term \u201cLand\u201d does not include any property beyond that described in Schedule A, nor any right, title, interest, estate, or easement in any abutting street, road, avenue, alley, lane, right-of-way, body of water, or waterway, but does not modify or limit the extent that a right of access to and from the Land is to be insured by the Policy."),
        ("d.", "\u201cMortgage\u201d: A mortgage, deed of trust, trust deed, security deed, or other real property security instrument, including one evidenced by electronic means authorized by law."),
        ("e.", "\u201cPolicy\u201d: Each contract of title insurance, in a form adopted by the American Land Title Association, issued or to be issued by the Company pursuant to this Commitment."),
        ("f.", "\u201cProposed Amount of Insurance\u201d: Each dollar amount specified in Schedule A as the Proposed Amount of Insurance of each Policy to be issued pursuant to this Commitment."),
        ("g.", "\u201cProposed Insured\u201d: Each person identified in Schedule A as the Proposed Insured of each Policy to be issued pursuant to this Commitment."),
        ("h.", "\u201cPublic Records\u201d: The recording or filing system established under State statutes in effect at the Commitment Date under which a document must be recorded or filed to impart constructive notice of matters relating to the Title to a purchaser for value without Knowledge. The term \u201cPublic Records\u201d does not include any other recording or filing system, including any pertaining to environmental remediation or protection, planning, permitting, zoning, licensing, building, health, public safety, or national security matters."),
        ("i.", "\u201cState\u201d: The state or commonwealth of the United States within whose exterior boundaries the Land is located. The term \u201cState\u201d also includes the District of Columbia, the Commonwealth of Puerto Rico, the U.S. Virgin Islands, and Guam."),
        ("j.", "\u201cTitle\u201d: The estate or interest in the Land identified in Item 3 of Schedule A."),
    ]
    for letter, text in defs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        r = p.add_run(f"{letter}\t{text}")
        r.font.size = Pt(10)

    numbered(doc, 2, "If all of the Schedule B, Part I\u2014Requirements have not been met within the time period specified in the Commitment to Issue Policy, this Commitment terminates and the Company\u2019s liability and obligation end.")

    p = numbered(doc, 3, "The Company\u2019s liability and obligation is limited by and this Commitment is not valid without:")
    for i, item in enumerate(["the Notice;", "the Commitment to Issue Policy;", "the Commitment Conditions;", "Schedule A;", "Schedule B, Part I\u2014Requirements;", "Schedule B, Part II\u2014Exceptions; and", "a counter-signature by the Company or its issuing agent that may be in electronic form."]):
        sp = doc.add_paragraph()
        sp.paragraph_format.left_indent = Inches(0.5)
        sp.add_run(f"{chr(97+i)}.\t{item}")

    p = doc.add_paragraph()
    r = p.add_run("4.\tCOMPANY\u2019S RIGHT TO AMEND")
    r.bold = True
    doc.add_paragraph(
        "The Company may amend this Commitment at any time. If the Company amends this Commitment to add "
        "a defect, lien, encumbrance, adverse claim, or other matter recorded in the Public Records prior to "
        "the Commitment Date, any liability of the Company is limited by Commitment Condition 5. The Company "
        "is not liable for any other amendment to this Commitment."
    )

    p = doc.add_paragraph()
    r = p.add_run("5.\tLIMITATIONS OF LIABILITY")
    r.bold = True
    for item in [
        "a.\tThe Company\u2019s liability under Commitment Condition 4 is limited to the Proposed Insured\u2019s actual expense incurred in the interval between the Company\u2019s delivery to the Proposed Insured of the Commitment and the delivery of the amended Commitment, resulting from the Proposed Insured\u2019s good faith reliance to:\n\ti.\tcomply with the Schedule B, Part I\u2014Requirements;\n\tii.\teliminate, with the Company\u2019s written consent, any Schedule B, Part II\u2014Exceptions; or\n\tiii.\tacquire the Title or create the Mortgage covered by this Commitment.",
        "b.\tThe Company is not liable under Commitment Condition 5.a. if the Proposed Insured requested the amendment or had Knowledge of the matter and did not notify the Company about it in writing.",
        "c.\tThe Company is only liable under Commitment Condition 4 if the Proposed Insured would not have incurred the expense had the Commitment included the added matter when the Commitment was first delivered to the Proposed Insured.",
        "d.\tThe Company\u2019s liability does not exceed the lesser of the Proposed Insured\u2019s actual expense incurred in good faith and described in Commitment Condition 5.a. or the Proposed Amount of Insurance.",
        "e.\tThe Company is not liable for the content of the Transaction Identification Data, if any.",
        "f.\tThe Company is not obligated to issue the Policy referred to in this Commitment unless all of the Schedule B, Part I\u2014Requirements have been met to the satisfaction of the Company.",
        "g.\tThe Company\u2019s liability is further limited by the terms and provisions of the Policy to be issued to the Proposed Insured.",
    ]:
        doc.add_paragraph(item)

    p = doc.add_paragraph()
    r = p.add_run("6.\tLIABILITY OF THE COMPANY MUST BE BASED ON THIS COMMITMENT; CHOICE OF LAW AND CHOICE OF FORUM")
    r.bold = True
    for item in [
        "a.\tOnly a Proposed Insured identified in Schedule A, and no other person, may make a claim under this Commitment.",
        "b.\tAny claim must be based in contract under the State law of the State where the Land is located and is restricted to the terms and provisions of this Commitment. Any litigation or other proceeding brought by the Proposed Insured against the Company must be filed only in a State or federal court having jurisdiction.",
        "c.\tThis Commitment, as last revised, is the exclusive and entire agreement between the parties with respect to the subject matter of this Commitment and supersedes all prior commitment negotiations, representations, and proposals of any kind, whether written or oral, express or implied, relating to the subject matter of this Commitment.",
        "d.\tThe deletion or modification of any Schedule B, Part II\u2014Exception does not constitute an agreement or obligation to provide coverage beyond the terms and provisions of this Commitment or the Policy.",
        "e.\tAny amendment or endorsement to this Commitment must be in writing and authenticated by a person authorized by the Company.",
        "f.\tWhen the Policy is issued, all liability and obligation under this Commitment will end and the Company\u2019s only liability will be under the Policy.",
    ]:
        doc.add_paragraph(item)

    for num, title, body in [
        (7, "IF THIS COMMITMENT IS ISSUED BY AN ISSUING AGENT", "The issuing agent is the Company\u2019s agent only for the limited purpose of issuing title insurance commitments and policies. The issuing agent is not the Company\u2019s agent for closing, settlement, escrow, or any other purpose."),
        (8, "PRO-FORMA POLICY", "The Company may provide, at the request of a Proposed Insured, a pro-forma policy illustrating the coverage that the Company may provide. A pro-forma policy neither reflects the status of Title at the time that the pro-forma policy is delivered to a Proposed Insured, nor is it a commitment to insure."),
        (9, "CLAIMS PROCEDURES", "This Commitment incorporates by reference all Conditions for making a claim in the Policy to be issued to the Proposed Insured. Commitment Condition 9 does not modify the limitations of liability in Commitment Conditions 5 and 6."),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"{num}.\t{title}")
        r.bold = True
        doc.add_paragraph(body)

    p = doc.add_paragraph()
    r = p.add_run("10.\tCLASS ACTION")
    r.bold = True
    doc.add_paragraph(
        "ALL CLAIMS AND DISPUTES ARISING OUT OF OR RELATING TO THIS COMMITMENT, INCLUDING ANY SERVICE OR "
        "OTHER MATTER IN CONNECTION WITH ISSUING THIS COMMITMENT, ANY BREACH OF A COMMITMENT PROVISION, OR "
        "ANY OTHER CLAIM OR DISPUTE ARISING OUT OF OR RELATING TO THE TRANSACTION GIVING RISE TO THIS "
        "COMMITMENT, MUST BE BROUGHT IN AN INDIVIDUAL CAPACITY. NO PARTY MAY SERVE AS PLAINTIFF, CLASS "
        "MEMBER, OR PARTICIPANT IN ANY CLASS OR REPRESENTATIVE PROCEEDING. ANY POLICY ISSUED PURSUANT TO "
        "THIS COMMITMENT WILL CONTAIN A CLASS ACTION CONDITION."
    )

    p = doc.add_paragraph()
    r = p.add_run("11.\tARBITRATION")
    r.bold = True
    doc.add_paragraph(
        "The Policy contains an arbitration clause. All arbitrable matters when the Proposed Amount of "
        "Insurance is $2,000,000 or less may be arbitrated at the election of either the Company or the "
        "Proposed Insured as the exclusive remedy of the parties. A Proposed Insured may review a copy of "
        "the arbitration rules at http://www.alta.org/arbitration."
    )

    add_footer(doc)
    doc.add_page_break()

    # ===================== PAGE 5 - TRANSACTION ID + SCHEDULE A =====================
    add_right_header(doc)

    p = doc.add_paragraph()
    r = p.add_run("Transaction Identification Data, for which the Company assumes no liability as set forth in Commitment Condition 5.e.:")
    r.bold = True

    for label, val in [
        ("Issuing Agent:", "[____________________]"),
        ("Issuing Office:", "[____________________]"),
        ("", "[____________________]"),
        ("Issuing Office\u2019s ALTA\u00ae Registry ID:", ""),
        ("Loan ID Number:", ""),
        ("Commitment Number:", "[____________________]"),
        ("Issuing Office File Number:", "[____________________]"),
        ("Property Address:", "236 Senior Road (F/K/A 512 Kinder Road), Kingstree, SC"),
        ("Revision Number:", ""),
    ]:
        p = doc.add_paragraph()
        if label:
            p.add_run(f"{label}  ")
        if val and '[' in val:
            hl(p, val)
        elif val:
            p.add_run(val)

    doc.add_paragraph()
    centered_bold(doc, "SCHEDULE A", 13)
    doc.add_paragraph()

    # 1
    p = numbered(doc, 1, "Commitment Date:  ")
    hl(p, "[____________________]")
    p.add_run(" at 8:00 AM")

    doc.add_paragraph()

    # 2
    numbered(doc, 2, "Policy to be issued:")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("a.\t2021 ALTA Homeowner\u2019s Policy\n")
    p.add_run("\tProposed Insured:\t\t")
    hl(p, "[____________________]")
    p.add_run("\n\tProposed Amount of Insurance:\t")
    hl(p, "[____________________]")
    p.add_run("\n\tThe estate or interest to be insured:\tfee simple")

    doc.add_paragraph()

    # 3
    numbered(doc, 3, "The estate or interest in the Land at the Commitment Date is:\n\tfee simple")
    doc.add_paragraph()

    # 4
    numbered(doc, 4,
        "The Title is, at the Commitment Date, vested in:\n"
        "\tChristopher Irvin, as Trustee of the John Earle Realty Trust, a trust dated May 19, 2009, "
        "by deed from Wade R.D. Sauls dated May 19, 2009 and recorded with the Williamsburg County "
        "Register of Deeds on June 26, 2009 in Book 120, Page 86 (Instrument #200900001862)."
    )
    doc.add_paragraph()

    # 5
    numbered(doc, 5, "The Land is described as follows:")
    doc.add_paragraph()
    centered_bold(doc, "SEE EXHIBIT A ATTACHED HERETO AND MADE A PART HEREOF", 11)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    hl(p, "[ISSUING AGENT NAME]").bold = True
    p = doc.add_paragraph()
    hl(p, "[Address]")
    doc.add_paragraph("Telephone:")
    doc.add_paragraph()
    doc.add_paragraph("Countersigned:")
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("By:____________________________")
    r.bold = True
    p = doc.add_paragraph()
    r = p.add_run("\tAuthorized Signatory")
    r.bold = True
    p = doc.add_paragraph()
    hl(p, "[Name, License #]")
    p = doc.add_paragraph()
    hl(p, "[Firm Name, License #]")

    add_footer(doc)
    doc.add_page_break()

    # ===================== PAGE 6 - EXHIBIT A =====================
    add_right_header(doc)
    centered_bold(doc, "EXHIBIT A", 13)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("Commitment No.: ")
    r.bold = True
    r.underline = True
    hl(p, "[____________________]")

    doc.add_paragraph()
    doc.add_paragraph(
        "The land referred to herein below is situated in the County of Williamsburg, State of South "
        "Carolina, and is described as follows:"
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "ALL that certain, parcel or lot of land, with the improvements thereon, situate in King "
        "Township, County of Williamsburg, State of South Carolina, in shape a rectangular parallelogram "
        "measuring two hundred ten (210) feet on its northern and southern boundary lines and one hundred "
        "and five (105) feet on its eastern and westerly boundary lines and bounded as follows: On the "
        "North by Lot Number XIV on the map hereinafter referred to, a lot hereinafter devised to Gerald "
        "M. Sauls; on the South by a twenty-eight (28) foot road separating the same from a lot heretofore "
        "conveyed to Reba S. Thornhill; and on the West by Kinder Street as shown on the map hereinafter "
        "referred to, being the northwestern portion of a two and one-half (2 1/2) acre tract of land "
        "known as designated as Lot Number XVI on a map showing subdivision of the lands of the Estate of "
        "W.H. Kinder made by J.D. Brockington and W.J. Green, Surveyors, in January 1953, and recorded in "
        "the Office of the Clerk of Court for Williamsburg County in Plat Book \"5\" at page 204, which said "
        "larger lot of land was conveyed to J.E. Sauls by S.L. Baylor by deed dated February 10, 1953 and "
        "recorded in the Office of the Clerk of Court for Williamsburg County in Deed Book \"A-51\" at page 174."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("TMS: 18-032-062")
    r.bold = True

    add_footer(doc)
    doc.add_page_break()

    # ===================== PAGE 7 - SCHEDULE B PART I =====================
    add_right_header(doc)
    centered_bold(doc, "SCHEDULE B, PART I\u2014Requirements", 13)
    doc.add_paragraph()
    doc.add_paragraph("All of the following Requirements must be met:")
    doc.add_paragraph()

    reqs = [
        "The Proposed Insured must notify the Company in writing of the name of any party not referred to in this Commitment who will obtain an interest in the Land or who will make a loan on the Land. The Company may then make additional Requirements or Exceptions.",
        "Pay the agreed amount for the estate or interest to be insured.",
        "Pay the premiums, fees, and charges for the Policy to the Company.",
        "We must be furnished with a copy of SCID 3601 executed pursuant to Section 38-75-960 S.C. Code of Laws 1976, as amended, and an executed Notice of Availability of Title Insurance pursuant to S.C. Insurance Department Regulation R-69-18, Vol. 25A of S.C. Code of Laws 1976, as amended.",
        "Seller\u2019s/Owner\u2019s Affidavit Indemnity executed by current owner(s) of the land on a form to be supplied by the Company stating that there have been no improvements to the land within the past 90 days which could give rise to a construction lien and that there are no accounts or claims pending and unpaid which could constitute a lien against the land. The affidavit will also state that affiant has no knowledge of any natural person or legal entity who has or could have a claim of right, interest or lien adverse to the Insured.",
        "Receipt of the acknowledged Privacy Policy.",
    ]
    for i, req in enumerate(reqs, 1):
        numbered(doc, i, req)

    # Requirement 7 - deed
    p = numbered(doc, 7, "Documents satisfactory to the Company that convey the Title or create the Mortgage to be insured, or both, must be properly authorized, executed, delivered, and recorded in the Public Records.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run("Warranty Deed from Christopher Irvin, as Trustee of the John Earle Realty Trust, to ")
    hl(p, "[BUYER NAME]")
    p.add_run(" conveying the land described in Schedule A herein.")

    add_footer(doc)
    doc.add_page_break()

    # ===================== PAGES 8-9 - SCHEDULE B PART II =====================
    add_right_header(doc)
    centered_bold(doc, "SCHEDULE B, PART II\u2014Exceptions", 13)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        "Some historical land records contain Discriminatory Covenants that are illegal and unenforceable by law. "
        "This Commitment and the Policy treat any Discriminatory Covenant in a document referenced in Schedule B "
        "as if each Discriminatory Covenant is redacted, repudiated, removed, and not republished or recirculated. "
        "Only the remaining provisions of the document will be excepted from coverage."
    )
    r.bold = True
    r.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        "The Policy will not insure against loss or damage resulting from the terms and conditions of any "
        "lease or easement identified in Schedule A, and will include the following Exceptions unless cleared "
        "to the satisfaction of the Company:"
    )
    doc.add_paragraph()

    # Standard exceptions 1-6
    std = [
        "Any defect, lien, encumbrance, adverse claim, or other matter that appears for the first time in the Public Records or is created, attaches, or is disclosed between the Commitment Date and the date on which all of the Schedule B, Part I\u2014Requirements are met.",
        "(a) Taxes or assessments that are not shown as existing liens by the records of any taxing authority that levies taxes or assessments on real property or by the Public Records; (b) proceedings by a public agency that may result in taxes or assessments, or notices of such proceedings, whether or not shown by the records of such agency or by the Public Records.",
        "Any facts, rights, interests, or claims that are not shown by the Public Records but that could be ascertained by an inspection of the Land or that may be asserted by persons in possession of the Land.",
        "Easements, liens or encumbrances, or claims thereof, not shown by the Public Records.",
        "Any encroachment, encumbrance, violation, variation, or adverse circumstance affecting the Title that would be disclosed by an accurate and complete land survey of the Land and not shown by the Public Records.",
        "Any mineral or mineral rights leased, granted or retained by current or prior owners.",
    ]
    for i, exc in enumerate(std, 1):
        numbered(doc, i, exc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("NOTE: Exceptions Numbered above will be hereby deleted upon issuance of the Loan Policy Only.")
    r.bold = True
    doc.add_paragraph()

    # Property-specific exceptions
    specific = [
        (7, "Taxes and assessments for the year 2026, and subsequent years, not yet due and payable."),
        (8, "Subject to a 28-foot road easement created by the Estate of J.E. Sauls along the southern boundary "
            "of the subject parcel, as set forth in Book 328 at Page 9 in the official records of the Williamsburg "
            "County Register of Deeds. Said easement impacts the subject parcel and adjacent lands formerly owned "
            "by the decedent, separating the subject parcel from the lot conveyed to Reba S. Thornhill, and "
            "providing access to interior lots created from the Sauls estate subdivision."),
        (9, "Subject to the Plat of the Estate of W.H. Kinder recorded in January 1953 in Plat Book 5 at "
            "Page 204 in the Williamsburg County Clerk of Court, showing the subdivision of the Kinder lands "
            "into lots, including the subject parcel being the northwestern portion of Lot XVI."),
        (10, "Subject to all rights, easements, and interests of the public and/or adjoining landowners in and "
             "to Kinder Street and the 28-foot road bounding the subject property."),
        (11, "The Company may make other requirements or exceptions upon its review of the proposed documents "
             "creating the estate or interest to be insured or otherwise ascertaining details of the transaction."),
    ]
    for num, text in specific:
        numbered(doc, num, text)

    add_footer(doc)

    # Save
    doc.save(OUTPUT)
    print(f"Done! Saved to: {OUTPUT}")


if __name__ == "__main__":
    build()
