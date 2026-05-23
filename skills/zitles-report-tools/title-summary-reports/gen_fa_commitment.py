#!/usr/bin/env python3
"""Generate a First American Title Insurance 2021 ALTA Commitment DOCX
that mimics the Long Grove form (Form 50133045 7-5-22).

Boilerplate + logo + "Page X of Y" live in the Word section header so they
repeat on every page and auto-update to the real page count.
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ASSETS = Path(__file__).parent / "assets"
FA_LOGO = ASSETS / "first_american_logo.jpeg"


# ---------- helpers ----------

def _set_font(run, *, name="Arial", size=10, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(parent, text, *, bold=False, italic=False, size=10, align=None,
             space_after=Pt(6), space_before=Pt(0), font="Arial"):
    p = parent.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, name=font, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    return p


def add_numbered(parent, number, text, *, size=10, font="Arial", space_after=Pt(6)):
    p = parent.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    r1 = p.add_run(f"{number}.\t"); _set_font(r1, name=font, size=size)
    r2 = p.add_run(text);            _set_font(r2, name=font, size=size)
    p.paragraph_format.space_after = space_after
    return p


def add_sub(parent, label, text, *, size=9, font="Arial", indent=0.6, space_after=Pt(3)):
    p = parent.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    r = p.add_run(f"{label}\t{text}")
    _set_font(r, name=font, size=size)
    p.paragraph_format.space_after = space_after
    return p


def add_field(paragraph, field_code, *, size=8, font="Arial"):
    """Insert a Word field (PAGE, NUMPAGES, etc.) into a paragraph."""
    run = paragraph.add_run()
    _set_font(run, name=font, size=size)

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "
    run._r.append(instr)

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_sep)

    # placeholder (word renders the real value when opened)
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    run._r.append(placeholder)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)


def _set_header(section):
    """Top-of-page header: logo at top-left, title right-aligned at top-right."""
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Two-column table: left = logo, right = title (right-aligned)
    tbl = header.add_table(rows=1, cols=2, width=Inches(6.7))
    tbl.autofit = False
    left, right = tbl.rows[0].cells
    left.width = Inches(3.0)
    right.width = Inches(3.7)

    # Remove default borders
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)

    # Clear default empty paragraphs in cells
    for cell in (left, right):
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)

    # Left cell: logo
    p_logo = left.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_after = Pt(0)
    if FA_LOGO.exists():
        p_logo.add_run().add_picture(str(FA_LOGO), width=Inches(2.0))

    # Right cell: title lines, right-aligned
    p_title = right.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_title.add_run("Commitment for Title Insurance")
    _set_font(r1, size=10, italic=True, bold=True)
    r2 = p_title.add_run("\nSouth Carolina - 2021 v. 01.00 (07-01-2021)")
    _set_font(r2, size=9, italic=True)
    p_title.paragraph_format.space_after = Pt(0)


def _set_footer(section):
    """Bottom-of-page footer: validity disclaimer + copyright + Form/Page."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    p_disc = footer.add_paragraph()
    r = p_disc.add_run(
        "This page is only a part of a 2021 ALTA Commitment for Title Insurance issued by First "
        "American Title Insurance Company. This Commitment is not valid without the Notice; the "
        "Commitment to Issue Policy; the Commitment Conditions; Schedule A; Schedule B, Part I\u2014"
        "Requirements; Schedule B, Part II\u2014Exceptions; and a counter-signature by the Company "
        "or its issuing agent that may be in electronic form."
    )
    _set_font(r, size=7)
    p_disc.paragraph_format.space_after = Pt(2)

    p_copy = footer.add_paragraph()
    r = p_copy.add_run(
        "Copyright 2021 American Land Title Association. All rights reserved.\n"
        "The use of this Form (or any derivative thereof) is restricted to ALTA licensees and "
        "ALTA members in good standing as of the date of use. All other uses are prohibited.\n"
        "Reprinted under license from the American Land Title Association."
    )
    _set_font(r, size=7, italic=True)
    p_copy.paragraph_format.space_after = Pt(4)

    p_form = footer.add_paragraph()
    p_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_form.add_run("Form 50133045 (7-5-22)    Page ")
    _set_font(r, size=9, italic=True)
    add_field(p_form, "PAGE", size=9)
    r = p_form.add_run(" of ")
    _set_font(r, size=9, italic=True)
    add_field(p_form, "NUMPAGES", size=9)
    p_form.paragraph_format.space_after = Pt(0)


def _force_update_fields(doc):
    """Set <w:updateFields w:val='true'/> in settings.xml so Word refreshes
    PAGE / NUMPAGES fields when the document is opened."""
    settings = doc.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def page_break(parent):
    p = parent.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------- content sections ----------

def section_notice_and_commitment(doc):
    add_para(doc, "ALTA COMMITMENT FOR TITLE INSURANCE", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(8), space_after=Pt(2))
    add_para(doc, "issued by", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_para(doc, "FIRST AMERICAN TITLE INSURANCE COMPANY", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_para(doc, "NOTICE", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc,
        "IMPORTANT\u2014READ CAREFULLY: THIS COMMITMENT IS AN OFFER TO ISSUE ONE OR MORE TITLE "
        "INSURANCE POLICIES. ALL CLAIMS OR REMEDIES SOUGHT AGAINST THE COMPANY INVOLVING THE "
        "CONTENT OF THIS COMMITMENT OR THE POLICY MUST BE BASED SOLELY IN CONTRACT.",
        size=8, bold=True)
    add_para(doc,
        "THIS COMMITMENT IS NOT AN ABSTRACT OF TITLE, REPORT OF THE CONDITION OF TITLE, LEGAL "
        "OPINION, OPINION OF TITLE, OR OTHER REPRESENTATION OF THE STATUS OF TITLE. THE "
        "PROCEDURES USED BY THE COMPANY TO DETERMINE INSURABILITY OF THE TITLE, INCLUDING ANY "
        "SEARCH AND EXAMINATION, ARE PROPRIETARY TO THE COMPANY, WERE PERFORMED SOLELY FOR "
        "THE BENEFIT OF THE COMPANY, AND CREATE NO EXTRACONTRACTUAL LIABILITY TO ANY PERSON, "
        "INCLUDING A PROPOSED INSURED.",
        size=8, bold=True)
    add_para(doc,
        "THE COMPANY\u2019S OBLIGATION UNDER THIS COMMITMENT IS TO ISSUE A POLICY TO A PROPOSED "
        "INSURED IDENTIFIED IN SCHEDULE A IN ACCORDANCE WITH THE TERMS AND PROVISIONS OF THIS "
        "COMMITMENT. THE COMPANY HAS NO LIABILITY OR OBLIGATION INVOLVING THE CONTENT OF THIS "
        "COMMITMENT TO ANY OTHER PERSON.",
        size=8, bold=True, space_after=Pt(10))

    add_para(doc, "COMMITMENT TO ISSUE POLICY", bold=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
        "Subject to the Notice; Schedule B, Part I\u2014Requirements; Schedule B, Part II\u2014Exceptions; "
        "and the Commitment Conditions, First American Title Insurance Company, a Nebraska Corporation "
        "(the \u201cCompany\u201d), commits to issue the Policy according to the terms and provisions of this "
        "Commitment. This Commitment is effective as of the Commitment Date shown in Schedule A for each "
        "Policy described in Schedule A, only when the Company has entered in Schedule A both the specified "
        "dollar amount as the Proposed Amount of Insurance and the name of the Proposed Insured.",
        size=9)
    add_para(doc,
        "If all of the Schedule B, Part I\u2014Requirements have not been met within 180 days after the "
        "Commitment Date, this Commitment terminates and the Company\u2019s liability and obligation end.",
        size=9, space_after=Pt(14))

    add_para(doc, "FIRST AMERICAN TITLE INSURANCE COMPANY", bold=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_para(doc, "By: ____________________________          By: ____________________________",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Sally F. Tyler, President                                    Lisa W. Cornehl, Secretary",
             size=8, align=WD_ALIGN_PARAGRAPH.CENTER)


def section_conditions(doc):
    page_break(doc)
    add_para(doc, "COMMITMENT CONDITIONS", bold=True, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    add_numbered(doc, 1, "DEFINITIONS")
    defs = [
        ("a.", "\u201cDiscriminatory Covenant\u201d: Any covenant, condition, restriction, or limitation that is unenforceable under applicable law because it illegally discriminates against a class of individuals based on personal characteristics such as race, color, religion, sex, sexual orientation, gender identity, familial status, disability, national origin, or other legally protected class."),
        ("b.", "\u201cKnowledge\u201d or \u201cKnown\u201d: Actual knowledge or actual notice, but not constructive notice imparted by the Public Records."),
        ("c.", "\u201cLand\u201d: The land described in Item 5 of Schedule A and improvements located on that land that by State law constitute real property. The term \u201cLand\u201d does not include any property beyond that described in Schedule A, nor any right, title, interest, estate, or easement in any abutting street, road, avenue, alley, lane, right-of-way, body of water, or waterway, but does not modify or limit the extent that a right of access to and from the Land is to be insured by the Policy."),
        ("d.", "\u201cMortgage\u201d: A mortgage, deed of trust, trust deed, security deed, or other real property security instrument, including one evidenced by electronic means authorized by law."),
        ("e.", "\u201cPolicy\u201d: Each contract of title insurance, in a form adopted by the American Land Title Association, issued or to be issued by the Company pursuant to this Commitment."),
        ("f.", "\u201cProposed Amount of Insurance\u201d: Each dollar amount specified in Schedule A as the Proposed Amount of Insurance of each Policy to be issued pursuant to this Commitment."),
        ("g.", "\u201cProposed Insured\u201d: Each person identified in Schedule A as the Proposed Insured of each Policy to be issued pursuant to this Commitment."),
        ("h.", "\u201cPublic Records\u201d: The recording or filing system established under State statutes in effect at the Commitment Date under which a document must be recorded or filed to impart constructive notice of matters relating to the Title to a purchaser for value without Knowledge. The term \u201cPublic Records\u201d does not include any other recording or filing system, including any pertaining to environmental remediation or protection, planning, permitting, zoning, licensing, building, health, public safety, or national security matters."),
        ("i.", "\u201cState\u201d: The state or commonwealth of the United States within whose exterior boundaries the Land is located. The term \u201cState\u201d also includes the District of Columbia, the Commonwealth of Puerto Rico, the U.S. Virgin Islands, and Guam."),
        ("j.", "\u201cTitle\u201d: The estate or interest in the Land identified in Item 3 of Schedule A."),
    ]
    for lbl, txt in defs:
        add_sub(doc, lbl, txt)

    add_numbered(doc, 2, "If all of the Schedule B, Part I\u2014Requirements have not been met within the time period specified in the Commitment to Issue Policy, this Commitment terminates and the Company\u2019s liability and obligation end.")
    add_numbered(doc, 3, "The Company\u2019s liability and obligation is limited by and this Commitment is not valid without:")
    for lbl, txt in [
        ("a.", "the Notice;"),
        ("b.", "the Commitment to Issue Policy;"),
        ("c.", "the Commitment Conditions;"),
        ("d.", "Schedule A;"),
        ("e.", "Schedule B, Part I\u2014Requirements;"),
        ("f.", "Schedule B, Part II\u2014Exceptions; and"),
        ("g.", "a counter-signature by the Company or its issuing agent that may be in electronic form."),
    ]:
        add_sub(doc, lbl, txt)

    add_numbered(doc, 4, "COMPANY\u2019S RIGHT TO AMEND")
    add_para(doc,
        "The Company may amend this Commitment at any time. If the Company amends this Commitment to add a "
        "defect, lien, encumbrance, adverse claim, or other matter recorded in the Public Records prior to "
        "the Commitment Date, any liability of the Company is limited by Commitment Condition 5. The Company "
        "is not liable for any other amendment to this Commitment.",
        size=9)

    add_numbered(doc, 5, "LIMITATIONS OF LIABILITY")
    for lbl, txt, indent in [
        ("a.", "The Company\u2019s liability under Commitment Condition 4 is limited to the Proposed Insured\u2019s actual expense incurred in the interval between the Company\u2019s delivery to the Proposed Insured of the Commitment and the delivery of the amended Commitment, resulting from the Proposed Insured\u2019s good faith reliance to:", 0.6),
        ("i.", "comply with the Schedule B, Part I\u2014Requirements;", 1.0),
        ("ii.", "eliminate, with the Company\u2019s written consent, any Schedule B, Part II\u2014Exceptions; or", 1.0),
        ("iii.", "acquire the Title or create the Mortgage covered by this Commitment.", 1.0),
        ("b.", "The Company is not liable under Commitment Condition 5.a. if the Proposed Insured requested the amendment or had Knowledge of the matter and did not notify the Company about it in writing.", 0.6),
        ("c.", "The Company is only liable under Commitment Condition 4 if the Proposed Insured would not have incurred the expense had the Commitment included the added matter when the Commitment was first delivered to the Proposed Insured.", 0.6),
        ("d.", "The Company\u2019s liability does not exceed the lesser of the Proposed Insured\u2019s actual expense incurred in good faith and described in Commitment Condition 5.a. or the Proposed Amount of Insurance.", 0.6),
        ("e.", "The Company is not liable for the content of the Transaction Identification Data, if any.", 0.6),
        ("f.", "The Company is not obligated to issue the Policy referred to in this Commitment unless all of the Schedule B, Part I\u2014Requirements have been met to the satisfaction of the Company.", 0.6),
        ("g.", "The Company\u2019s liability is further limited by the terms and provisions of the Policy to be issued to the Proposed Insured.", 0.6),
    ]:
        add_sub(doc, lbl, txt, indent=indent)

    add_numbered(doc, 6, "LIABILITY OF THE COMPANY MUST BE BASED ON THIS COMMITMENT; CHOICE OF LAW AND CHOICE OF FORUM")
    for lbl, txt in [
        ("a.", "Only a Proposed Insured identified in Schedule A, and no other person, may make a claim under this Commitment."),
        ("b.", "Any claim must be based in contract under the State law of the State where the Land is located and is restricted to the terms and provisions of this Commitment. Any litigation or other proceeding brought by the Proposed Insured against the Company must be filed only in a State or federal court having jurisdiction."),
        ("c.", "This Commitment, as last revised, is the exclusive and entire agreement between the parties with respect to the subject matter of this Commitment and supersedes all prior commitment negotiations, representations, and proposals of any kind, whether written or oral, express or implied, relating to the subject matter of this Commitment."),
        ("d.", "The deletion or modification of any Schedule B, Part II\u2014Exception does not constitute an agreement or obligation to provide coverage beyond the terms and provisions of this Commitment or the Policy."),
        ("e.", "Any amendment or endorsement to this Commitment must be in writing and authenticated by a person authorized by the Company."),
        ("f.", "When the Policy is issued, all liability and obligation under this Commitment will end and the Company\u2019s only liability will be under the Policy."),
    ]:
        add_sub(doc, lbl, txt)

    add_numbered(doc, 7, "IF THIS COMMITMENT IS ISSUED BY AN ISSUING AGENT")
    add_para(doc,
        "The issuing agent is the Company\u2019s agent only for the limited purpose of issuing title insurance "
        "commitments and policies. The issuing agent is not the Company\u2019s agent for closing, settlement, "
        "escrow, or any other purpose.",
        size=9)

    add_numbered(doc, 8, "PRO-FORMA POLICY")
    add_para(doc,
        "The Company may provide, at the request of a Proposed Insured, a pro-forma policy illustrating the "
        "coverage that the Company may provide. A pro-forma policy neither reflects the status of Title at the "
        "time that the pro-forma policy is delivered to a Proposed Insured, nor is it a commitment to insure.",
        size=9)

    add_numbered(doc, 9, "CLAIMS PROCEDURES")
    add_para(doc,
        "This Commitment incorporates by reference all Conditions for making a claim in the Policy to be issued "
        "to the Proposed Insured. Commitment Condition 9 does not modify the limitations of liability in "
        "Commitment Conditions 5 and 6.",
        size=9)

    add_numbered(doc, 10, "CLASS ACTION")
    add_para(doc,
        "ALL CLAIMS AND DISPUTES ARISING OUT OF OR RELATING TO THIS COMMITMENT, INCLUDING ANY SERVICE OR OTHER "
        "MATTER IN CONNECTION WITH ISSUING THIS COMMITMENT, ANY BREACH OF A COMMITMENT PROVISION, OR ANY OTHER "
        "CLAIM OR DISPUTE ARISING OUT OF OR RELATING TO THE TRANSACTION GIVING RISE TO THIS COMMITMENT, MUST BE "
        "BROUGHT IN AN INDIVIDUAL CAPACITY. NO PARTY MAY SERVE AS PLAINTIFF, CLASS MEMBER, OR PARTICIPANT IN ANY "
        "CLASS OR REPRESENTATIVE PROCEEDING. ANY POLICY ISSUED PURSUANT TO THIS COMMITMENT WILL CONTAIN A CLASS "
        "ACTION CONDITION.",
        size=9, bold=True)

    add_numbered(doc, 11, "ARBITRATION")
    add_para(doc,
        "The Policy contains an arbitration clause. All arbitrable matters when the Proposed Amount of Insurance "
        "is $2,000,000 or less may be arbitrated at the election of either the Company or the Proposed Insured as "
        "the exclusive remedy of the parties. A Proposed Insured may review a copy of the arbitration rules at "
        "http://www.alta.org/arbitration.",
        size=9)


def section_schedule_a(doc, d):
    page_break(doc)
    add_para(doc,
        "Transaction Identification Data, for which the Company assumes no liability as set forth in "
        "Commitment Condition 5.e.:", bold=True, size=9)
    for label, val in [
        ("Issuing Agent:", d.get("issuing_agent", "")),
        ("Issuing Office:", d.get("issuing_office", "")),
        ("Issuing Office\u2019s ALTA\u00ae Registry ID:", d.get("alta_registry_id", "")),
        ("Loan ID Number:", d.get("loan_id", "")),
        ("Commitment Number:", d.get("commitment_number", "")),
        ("Issuing Office File Number:", d.get("file_number", "")),
        ("Property Address:", d.get("property_address_full", "")),
        ("Revision Number:", d.get("revision_number", "")),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"{label} {val}"); _set_font(r, size=9)
        p.paragraph_format.space_after = Pt(2)

    add_para(doc, "SCHEDULE A", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(8), space_after=Pt(6))

    add_numbered(doc, 1, f"Commitment Date: {d.get('commitment_date', '')} at 8:00 AM")
    add_numbered(doc, 2, "Policy to be issued:")
    for lbl, txt in [
        ("a.", d.get("policy_type", "2021 ALTA Owner\u2019s Policy")),
        ("",   f"Proposed Insured: {d.get('proposed_insured', '_______________')}"),
        ("",   f"Proposed Amount of Insurance: ${d.get('proposed_amount', '_______________')}"),
        ("",   f"The estate or interest to be insured: {d.get('estate', 'fee simple')}"),
    ]:
        add_sub(doc, lbl, txt, size=10)

    add_numbered(doc, 3, f"The estate or interest in the Land at the Commitment Date is: {d.get('estate', 'fee simple')}")
    add_numbered(doc, 4, "The Title is, at the Commitment Date, vested in:")
    add_para(doc, d.get("vesting_paragraph", ""), size=10, space_after=Pt(6))
    if d.get("prior_chain_paragraph"):
        add_para(doc, d["prior_chain_paragraph"], size=10, space_after=Pt(6))

    add_numbered(doc, 5, "The Land is described as follows:")
    add_para(doc, "SEE EXHIBIT A ATTACHED HERETO AND MADE A PART HEREOF", bold=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(4), space_after=Pt(18))

    add_para(doc, d.get("issuing_agent_upper", "[ISSUING AGENT]"), bold=True, size=10)
    add_para(doc, d.get("issuing_office", ""), size=9)
    add_para(doc, f"Telephone: {d.get('agent_phone', '')}", size=9)
    add_para(doc, "Countersigned:", size=9, space_after=Pt(6))
    add_para(doc, "By: ____________________________", size=10)
    add_para(doc, "Authorized Signatory", size=9)
    if d.get("agent_license"):
        add_para(doc, d["agent_license"], size=9)


def section_exhibit_a(doc, d):
    page_break(doc)
    add_para(doc, "EXHIBIT A", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_para(doc, f"Commitment No.: {d.get('commitment_number', '')}", size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_para(doc,
        f"The land referred to herein below is situated in the County of {d.get('county', '')}, "
        "State of South Carolina, and is described as follows:",
        size=10, space_after=Pt(8))
    for para in d.get("legal_description_paragraphs", []):
        add_para(doc, para, size=10, space_after=Pt(6))
    if d.get("parcel_id"):
        add_para(doc, f"TMS / PIN: {d['parcel_id']}", size=10, space_before=Pt(6))


def section_schedule_b1(doc, d):
    page_break(doc)
    add_para(doc, "SCHEDULE B, PART I\u2014Requirements", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_para(doc, "All of the following Requirements must be met:", size=10, space_after=Pt(6))
    for i, req in enumerate(d.get("requirements", []), 1):
        add_numbered(doc, i, req, size=10)


def section_schedule_b2(doc, d):
    page_break(doc)
    add_para(doc, "SCHEDULE B, PART II\u2014Exceptions", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
    add_para(doc,
        "Some historical land records contain Discriminatory Covenants that are illegal and unenforceable by law. "
        "This Commitment and the Policy treat any Discriminatory Covenant in a document referenced in Schedule B "
        "as if each Discriminatory Covenant is redacted, repudiated, removed, and not republished or recirculated. "
        "Only the remaining provisions of the document will be excepted from coverage.",
        size=9, italic=True, space_after=Pt(6))
    add_para(doc,
        "The Policy will not insure against loss or damage resulting from the terms and conditions of any lease "
        "or easement identified in Schedule A, and will include the following Exceptions unless cleared to the "
        "satisfaction of the Company:",
        size=10, space_after=Pt(6))
    for i, exc in enumerate(d.get("exceptions", []), 1):
        add_numbered(doc, i, exc, size=10)


# ---------- main ----------

def build(data: dict, out_path: str):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(1.2)     # header = logo + title only
        section.bottom_margin = Inches(1.8)  # footer = disclaimer + copyright + form/page
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
        _set_header(section)
        _set_footer(section)

    section_notice_and_commitment(doc)
    section_conditions(doc)
    section_schedule_a(doc, data)
    section_exhibit_a(doc, data)
    section_schedule_b1(doc, data)
    section_schedule_b2(doc, data)

    _force_update_fields(doc)
    doc.save(out_path)
    print(f"Done! DOCX: {out_path}")


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("data_file")
    ap.add_argument("-o", "--output", required=True)
    args = ap.parse_args()
    with open(args.data_file) as f:
        data = json.load(f)
    build(data, args.output)
